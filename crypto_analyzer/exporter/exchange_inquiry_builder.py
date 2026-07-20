"""
exchange_inquiry_builder.py
加密貨幣交易所調閱案件申請單生成器
格式依據：加密貨幣交易所調閱案件申請單1140111.docx（OKX 官方範本）
支援格式：DOCX（python-docx）/ ODT（odfpy）/ PDF（reportlab）
"""
from __future__ import annotations
import os
import datetime
from typing import Any


def _find_default_logo() -> str:
    """從 exporter/ 資料夾往上一層尋找 logo1.jpg"""
    module_dir = os.path.dirname(os.path.abspath(__file__))
    candidate = os.path.normpath(os.path.join(module_dir, "..", "logo1.jpg"))
    return candidate if os.path.isfile(candidate) else ""


# ── 機關英文名稱靜態對照表 ──────────────────────────────────────────────────
# 來源：各機關官方公文英文名稱，優先於線上翻譯

AGENCY_EN_MAP: dict[str, str] = {
    # ── 檢察署 ──
    "臺灣臺北地方檢察署":          "Taipei District Prosecutors Office",
    "台灣台北地方檢察署":           "Taipei District Prosecutors Office",
    "臺灣士林地方檢察署":          "Shilin District Prosecutors Office",
    "台灣士林地方檢察署":          "Shilin District Prosecutors Office",
    "臺灣新北地方檢察署":          "New Taipei District Prosecutors Office",
    "台灣新北地方檢察署":          "New Taipei District Prosecutors Office",
    "臺灣板橋地方檢察署":          "Banqiao District Prosecutors Office",
    "台灣板橋地方檢察署":          "Banqiao District Prosecutors Office",
    "臺灣桃園地方檢察署":          "Taoyuan District Prosecutors Office",
    "台灣桃園地方檢察署":          "Taoyuan District Prosecutors Office",
    "臺灣新竹地方檢察署":          "Hsinchu District Prosecutors Office",
    "台灣新竹地方檢察署":          "Hsinchu District Prosecutors Office",
    "臺灣苗栗地方檢察署":          "Miaoli District Prosecutors Office",
    "台灣苗栗地方檢察署":          "Miaoli District Prosecutors Office",
    "臺灣臺中地方檢察署":          "Taichung District Prosecutors Office",
    "台灣台中地方檢察署":          "Taichung District Prosecutors Office",
    "臺灣南投地方檢察署":          "Nantou District Prosecutors Office",
    "台灣南投地方檢察署":          "Nantou District Prosecutors Office",
    "臺灣彰化地方檢察署":          "Changhua District Prosecutors Office",
    "台灣彰化地方檢察署":          "Changhua District Prosecutors Office",
    "臺灣雲林地方檢察署":          "Yunlin District Prosecutors Office",
    "台灣雲林地方檢察署":          "Yunlin District Prosecutors Office",
    "臺灣嘉義地方檢察署":          "Chiayi District Prosecutors Office",
    "台灣嘉義地方檢察署":          "Chiayi District Prosecutors Office",
    "臺灣臺南地方檢察署":          "Tainan District Prosecutors Office",
    "台灣台南地方檢察署":          "Tainan District Prosecutors Office",
    "臺灣高雄地方檢察署":          "Kaohsiung District Prosecutors Office",
    "台灣高雄地方檢察署":          "Kaohsiung District Prosecutors Office",
    "臺灣屏東地方檢察署":          "Pingtung District Prosecutors Office",
    "台灣屏東地方檢察署":          "Pingtung District Prosecutors Office",
    "臺灣花蓮地方檢察署":          "Hualien District Prosecutors Office",
    "台灣花蓮地方檢察署":          "Hualien District Prosecutors Office",
    "臺灣臺東地方檢察署":          "Taitung District Prosecutors Office",
    "台灣台東地方檢察署":          "Taitung District Prosecutors Office",
    "臺灣宜蘭地方檢察署":          "Yilan District Prosecutors Office",
    "台灣宜蘭地方檢察署":          "Yilan District Prosecutors Office",
    "臺灣基隆地方檢察署":          "Keelung District Prosecutors Office",
    "台灣基隆地方檢察署":          "Keelung District Prosecutors Office",
    "臺灣澎湖地方檢察署":          "Penghu District Prosecutors Office",
    "台灣澎湖地方檢察署":          "Penghu District Prosecutors Office",
    "臺灣金門地方檢察署":          "Kinmen District Prosecutors Office",
    "台灣金門地方檢察署":          "Kinmen District Prosecutors Office",
    "臺灣連江地方檢察署":          "Lienchiang District Prosecutors Office",
    "台灣連江地方檢察署":          "Lienchiang District Prosecutors Office",
    # 高等 / 最高
    "臺灣高等檢察署":              "Taiwan High Prosecutors Office",
    "台灣高等檢察署":              "Taiwan High Prosecutors Office",
    "臺灣高等檢察署臺中檢察分署":  "Taichung Branch, Taiwan High Prosecutors Office",
    "臺灣高等檢察署台中檢察分署":  "Taichung Branch, Taiwan High Prosecutors Office",
    "臺灣高等檢察署臺南檢察分署":  "Tainan Branch, Taiwan High Prosecutors Office",
    "臺灣高等檢察署台南檢察分署":  "Tainan Branch, Taiwan High Prosecutors Office",
    "臺灣高等檢察署高雄檢察分署":  "Kaohsiung Branch, Taiwan High Prosecutors Office",
    "最高檢察署":                  "Supreme Prosecutors Office",

    # ── 警察局 ──
    "內政部警政署刑事警察局":       "Criminal Investigation Bureau, National Police Agency, Ministry of the Interior",
    "刑事警察局":                  "Criminal Investigation Bureau (CIB)",
    "刑事局":                      "Criminal Investigation Bureau (CIB)",
    "法務部調查局":                 "Investigation Bureau, Ministry of Justice (MJIB)",
    "調查局":                      "Investigation Bureau, Ministry of Justice (MJIB)",
    "內政部警政署航空警察局":       "Airport Police Bureau, National Police Agency, Ministry of the Interior",
    "航空警察局":                   "Airport Police Bureau",
    "內政部警政署國道公路警察局":   "National Freeway Police Bureau, National Police Agency, Ministry of the Interior",
    "國道公路警察局":               "National Freeway Police Bureau",
    "臺北市政府警察局":             "Taipei City Police Department",
    "台北市政府警察局":             "Taipei City Police Department",
    "新北市政府警察局":             "New Taipei City Police Department",
    "桃園市政府警察局":             "Taoyuan City Police Department",
    "臺中市政府警察局":             "Taichung City Police Department",
    "台中市政府警察局":             "Taichung City Police Department",
    "臺南市政府警察局":             "Tainan City Police Department",
    "台南市政府警察局":             "Tainan City Police Department",
    "高雄市政府警察局":             "Kaohsiung City Police Department",
    "基隆市政府警察局":             "Keelung City Police Department",
    "新竹市政府警察局":             "Hsinchu City Police Department",
    "嘉義市政府警察局":             "Chiayi City Police Department",
    "宜蘭縣政府警察局":             "Yilan County Police Bureau",
    "苗栗縣政府警察局":             "Miaoli County Police Bureau",
    "彰化縣政府警察局":             "Changhua County Police Bureau",
    "南投縣政府警察局":             "Nantou County Police Bureau",
    "雲林縣政府警察局":             "Yunlin County Police Bureau",
    "嘉義縣政府警察局":             "Chiayi County Police Bureau",
    "屏東縣政府警察局":             "Pingtung County Police Bureau",
    "花蓮縣政府警察局":             "Hualien County Police Bureau",
    "臺東縣政府警察局":             "Taitung County Police Bureau",
    "台東縣政府警察局":             "Taitung County Police Bureau",
    "澎湖縣政府警察局":             "Penghu County Police Bureau",
    "金門縣政府警察局":             "Kinmen County Police Bureau",
    "連江縣政府警察局":             "Lienchiang County Police Bureau",
    "新竹縣政府警察局":             "Hsinchu County Police Bureau",

    # ── 海巡 / 憲兵 ──
    "行政院海洋委員會海巡署":       "Coast Guard Administration, Ocean Affairs Council",
    "海巡署":                      "Coast Guard Administration",
    "憲兵指揮部":                  "Military Police Command",
}


# ── 分局 / 派出所等單位類型對照 ──────────────────────────────────────────────
_UNIT_SUFFIX_MAP: dict[str, str] = {
    "分局":            "Precinct",
    "第一分局":        "1st Precinct",
    "第二分局":        "2nd Precinct",
    "第三分局":        "3rd Precinct",
    "派出所":          "Police Station",
    "刑事偵查大隊":    "Criminal Investigation Brigade",
    "刑事偵查隊":      "Criminal Investigation Squad",
    "刑事偵查組":      "Criminal Investigation Team",
    "偵查隊":          "Investigation Squad",
    "偵查組":          "Investigation Team",
    "交通隊":          "Traffic Police Squad",
    "保安隊":          "Security Squad",
    "婦幼隊":          "Women and Children Protection Team",
    "少年隊":          "Juvenile Investigation Team",
    "科技犯罪偵查隊":  "Cyber Crime Investigation Squad",
    "組織犯罪偵查隊":  "Organized Crime Investigation Squad",
}

# 「第一」「第二」等序數轉換（分局序號）
_ORDINAL_MAP: dict[str, str] = {
    "第一": "1st", "第二": "2nd", "第三": "3rd", "第四": "4th",
    "第五": "5th", "第六": "6th",
    "一": "1st", "二": "2nd", "三": "3rd",
}

# ── 縣市行政區地名拼音表 ──────────────────────────────────────────────────────
_DISTRICT_MAP: dict[str, str] = {
    # 臺北市
    "松山": "Songshan", "信義": "Xinyi", "大安": "Da'an", "萬華": "Wanhua",
    "中山": "Zhongshan", "大同": "Datong", "內湖": "Neihu", "南港": "Nangang",
    "士林": "Shilin", "北投": "Beitou", "中正": "Zhongzheng",
    "中正一": "Zhongzheng 1st", "中正二": "Zhongzheng 2nd", "文山": "Wenshan",
    # 新北市
    "板橋": "Banqiao", "三重": "Sanchong", "新莊": "Xinzhuang", "中和": "Zhonghe",
    "永和": "Yonghe", "土城": "Tucheng", "樹林": "Shulin", "鶯歌": "Yingge",
    "三峽": "Sanxia", "淡水": "Tamsui", "汐止": "Xizhi", "瑞芳": "Ruifang",
    "新店": "Xindian", "蘆洲": "Luzhou", "五股": "Wugu", "泰山": "Taishan",
    "林口": "Linkou", "八里": "Bali", "金山": "Jinshan", "萬里": "Wanli",
    # 桃園市
    "桃園": "Taoyuan", "中壢": "Zhongli", "平鎮": "Pingzhen", "八德": "Bade",
    "楊梅": "Yangmei", "龜山": "Guishan", "龍潭": "Longtan", "大溪": "Daxi",
    "大園": "Dayuan", "觀音": "Guanyin", "蘆竹": "Luzhu",
    # 新竹
    "新竹": "Hsinchu", "竹北": "Zhubei", "竹東": "Zhudong",
    # 臺中市
    "豐原": "Fengyuan", "大里": "Dali", "太平": "Taiping", "清水": "Qingshui",
    "沙鹿": "Shalu", "梧棲": "Wuqi", "烏日": "Wuri", "大甲": "Dajia",
    "后里": "Houli", "東勢": "Dongshi", "霧峰": "Wufeng", "潭子": "Tanzi",
    # 臺南市
    "永康": "Yongkang", "仁德": "Rende", "歸仁": "Guiren", "新營": "Sinying",
    "麻豆": "Madou", "佳里": "Jiali", "安南": "Annan", "安平": "Anping",
    "東山": "Dongshan", "白河": "Baihe",
    # 高雄市
    "三民": "Sanmin", "三民第一": "Sanmin 1st", "三民第二": "Sanmin 2nd",
    "鹽埕": "Yancheng", "鼓山": "Gushan", "前金": "Qianjin",
    "苓雅": "Lingya", "前鎮": "Qianzhen", "小港": "Xiaogang",
    "楠梓": "Nanzi", "左營": "Zuoying", "仁武": "Renwu", "鳳山": "Fengshan",
    "岡山": "Gangshan", "路竹": "Luzhu", "橋頭": "Qiaotou", "大寮": "Daliao",
    "林園": "Linyuan", "旗山": "Qishan", "美濃": "Meinong",
    # 屏東 / 花蓮 / 臺東 / 其他
    "屏東": "Pingtung", "潮州": "Chaozhou", "恆春": "Hengchun",
    "花蓮": "Hualien", "吉安": "Ji'an",
    "臺東": "Taitung", "台東": "Taitung",
    "宜蘭": "Yilan", "羅東": "Luodong", "蘇澳": "Su'ao",
    "基隆": "Keelung", "七堵": "Qidu", "中山": "Zhongshan",
}


def _parse_suffix_en(suffix: str) -> str | None:
    """
    嘗試將單位尾碼（如「松山分局」「三民第一分局」）解析為英文。
    策略：先比對完整 _UNIT_SUFFIX_MAP，再拆分地名 + 單位類型。
    """
    # 1. 直接比對完整尾碼（如「第一分局」）
    if suffix in _UNIT_SUFFIX_MAP:
        return _UNIT_SUFFIX_MAP[suffix]

    # 2. 找最長匹配的單位類型（suffix 結尾）
    for unit_cn, unit_en in sorted(_UNIT_SUFFIX_MAP.items(),
                                    key=lambda x: -len(x[0])):
        if suffix.endswith(unit_cn):
            location = suffix[:-len(unit_cn)].strip()
            if not location:
                return unit_en
            # 地名查表
            if location in _DISTRICT_MAP:
                return f"{_DISTRICT_MAP[location]} {unit_en}"
            # 地名含序數（如「中正一」「三民第二」）
            for ord_cn, ord_en in _ORDINAL_MAP.items():
                if location.endswith(ord_cn):
                    base_loc = location[:-len(ord_cn)]
                    if base_loc in _DISTRICT_MAP:
                        return f"{_DISTRICT_MAP[base_loc]} {ord_en} {unit_en}"
            # 地名本身也是複合尾碼（如「松山分局」）→ 遞迴解析
            # 例：「松山分局刑事偵查隊」→ unit="刑事偵查隊", location="松山分局"
            #     → 遞迴 _parse_suffix_en("松山分局") = "Songshan Precinct"
            #     → 結果 "Criminal Investigation Squad, Songshan Precinct"
            sub_en = _parse_suffix_en(location)
            if sub_en:
                return f"{unit_en}, {sub_en}"
            # 地名查無 → 回傳 None，讓呼叫方降級到 API
            return None
    return None


def get_agency_en(cn_name: str) -> str | None:
    """
    精確比對對照表。只做完整名稱查詢，不處理尾碼。
    """
    return AGENCY_EN_MAP.get(cn_name.strip())


def translate_via_api(text: str, src: str = "zh-TW", tgt: str = "en",
                      timeout: int = 8) -> str | None:
    """
    呼叫 MyMemory 免費翻譯 API（無需 API Key，每日 5000 字元）。
    失敗時回傳 None。
    """
    try:
        import urllib.request
        import urllib.parse
        import json
        params = urllib.parse.urlencode({"q": text, "langpair": f"{src}|{tgt}"})
        url = f"https://api.mymemory.translated.net/get?{params}"
        with urllib.request.urlopen(url, timeout=timeout) as resp:
            data = json.loads(resp.read().decode())
        if data.get("responseStatus") == 200:
            result = data["responseData"]["translatedText"]
            # MyMemory 在配額超出時回傳錯誤字串
            if "MYMEMORY WARNING" not in result.upper():
                return result
    except Exception:
        pass
    return None


def auto_translate_agency(cn_name: str) -> tuple[str, str]:
    """
    回傳 (英文名稱, 來源)。
    來源值：'table'（完整精確）/ 'table+api'（基礎機關+尾碼 API）/
            'table+partial'（API 失敗，尾碼保留原文）/ 'api'（全名 API）/ 'fail'

    處理流程：
      1. 精確比對 → 直接回傳
      2. 前綴比對（找最長基礎機關，取剩餘尾碼）→ 尾碼送 API 翻譯 → 組合
         格式：「{尾碼英文}, {基礎機關英文}」
         例：臺北市政府警察局松山分局 → Songshan Precinct, Taipei City Police Department
      3. 全名送 API
    """
    name = cn_name.strip()

    # 1. 精確比對
    if name in AGENCY_EN_MAP:
        return AGENCY_EN_MAP[name], "table"

    # 2. 前綴比對：找出對照表中以 name 為前綴的最長鍵
    best_key, best_val = "", ""
    for key, val in AGENCY_EN_MAP.items():
        if name.startswith(key) and len(key) > len(best_key):
            best_key, best_val = key, val

    if best_val:
        suffix = name[len(best_key):].strip()
        if not suffix:
            return best_val, "table"
        # 有尾碼（如「松山分局」「刑事偵查隊」）
        # 優先用本地表解析，格式：{尾碼英文}, {基礎機關英文}
        suffix_en = _parse_suffix_en(suffix)
        if suffix_en:
            return f"{suffix_en}, {best_val}", "table"
        # 本地表查無 → 呼叫 API 翻譯尾碼
        suffix_en = translate_via_api(suffix)
        if suffix_en:
            return f"{suffix_en}, {best_val}", "table+api"
        # API 也失敗 → 保留原文尾碼提示
        return f"{best_val} ({suffix})", "table+partial"

    # 3. 完全無法比對 → 全名送 API
    en = translate_via_api(name)
    if en:
        return en, "api"
    return "", "fail"


def translate_long(text: str, chunk_size: int = 450) -> str | None:
    """
    將長文字分段翻譯（MyMemory 每次上限 ~500 字元）。
    依句號「。」「！」「？」「；」分段，每段不超過 chunk_size 字元。
    回傳合併後的英文字串，或 None（全部失敗）。
    """
    import re
    # 依中文標點切句
    sentences = re.split(r'(?<=[。！？；\n])', text)
    chunks: list[str] = []
    current = ""
    for s in sentences:
        if len(current) + len(s) <= chunk_size:
            current += s
        else:
            if current:
                chunks.append(current)
            current = s
    if current:
        chunks.append(current)

    parts: list[str] = []
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        en = translate_via_api(chunk)
        if en is None:
            return None
        parts.append(en)
    return " ".join(parts) if parts else None


def auto_translate_address(cn_addr: str) -> tuple[str, str]:
    """
    地址使用 MyMemory API 翻譯（無靜態表）。
    回傳 (英文地址, 來源)
    """
    en = translate_via_api(cn_addr)
    if en:
        return en, "api"
    return "", "fail"


# ── 已知交易所與聯絡 Email ───────────────────────────────────────────────────

KNOWN_EXCHANGES: dict[str, str] = {
    "OKX":       "enforcement@okx.com",
    "Binance":   "law.enforcement@binance.com",
    "Bitget":    "compliance@bitget.com",
    "HTX":       "compliance@htx.com",
    "Bybit":     "compliance@bybit.com",
    "KuCoin":    "compliance@kucoin.com",
    "Gate.io":   "regulatory@gate.com",
    "Kraken":    "lawenforcement@kraken.com",
    "Coinbase":  "records@coinbase.com",

    # ── 以下取自 D:\交易所調閱資訊\ 內 48 份 Kodex Global 政府機構申請入口
    # 網站說明文件（彙整於 D:\claude\交易所調閱資訊彙整表.md，2026-07-10）
    # email 欄位留空者，代表該交易所僅提供 Kodex 入口網站/線上表單受理，
    # 無公開直接 email，送件前請至彙整表對應之入口網址查詢或提交
    "1inch":                      "",
    "bingX":                      "",
    "Bitbuy":                     "",
    "Coinbase APAC":              "",
    "Coinsquare":                 "",
    "Coin Wallet":                "",
    "Darkex":                     "",
    "Exodus":                     "",
    "Kraken-Rest of World":       "",
    "MaskEx":                     "",
    "NEAR Intents":               "",
    "Rain":                       "",
    "Strike":                     "",
    "Tether":                     "",
    "Tools For Humanity":         "",
    "Athena Bitcoin":             "subpoenas@athenabitcoin.com",
    "BigONE":                     "fraud-report@bigone.com",
    "Bitcoin Depot":              "lawenforcement@bitcoindepot.com",
    "WhiteBIT":                   "legal@whitebit.com",
    "Paxful":                     "legal@paxful.com",
    "Nexo Group":                 "compliance@nexo.io",
    "MEXC":                       "",
    "Magic Eden":                 "legal@magiceden.io",
    "HitBTC":                     "legal@hitbtc.com",
    "Gemini":                     "lawenforcement@gemini.com",
    "FixedFloat":                 "info@fixedfloat.com",
    "CoinsPaid":                  "compliance@coinspaid.com",
    "Coinme":                     "comp-sec@coinme.com",
    "Coinhub":                    "compliance@coinhubatm.com",
    "CoinGate":                   "",
    "CoinCola":                   "evidence@coincola.com",
    "Cobo":                       "amlcompliance@cobo.com",
    "ChangeNOW":                  "compliance@changenow.io",
    "Changelly":                  "legal@changelly.com",
    "Bittrex":                    "LErequests@bittrex.com",
    "Bitstamp":                   "compliance@bitstamp.net",
    "Bitso":                      "Requests@bitso.com",
    "Bitrefill":                  "legal@bitrefill.com",
    "Bitpay":                     "subpoenas@bitpay.com",
    "Crypto.com (Rest of World)": "",

    "其他":      "",
}

# ── 案件性質（依範本順序與分行排列） ─────────────────────────────────────────
# 每個 tuple：(中文名稱, 英文名稱)
# 前 6 項三個一排，後續每項一行

CASE_TYPES: list[tuple[str, str]] = [
    ("殺人罪",             "Homicide"),
    ("擄人勒贖罪",         "Kidnapping"),
    ("毒品罪",             "Drug Offence"),
    ("詐欺罪",             "Scams and Frauds"),
    ("竊盜罪",             "Theft"),
    ("強盜罪",             "Robbery"),
    ("搶奪罪",             "Snatch theft"),
    ("違反組織犯罪防制條例", "Organized Crime"),
    ("違反槍砲彈藥刀械管制條例", "Violation of Gun Control Act"),
    ("傷害罪(含重傷害)",   "Offenses of Causing Bodily Harm"),
    ("妨害性自主",         "Sexual Assault"),
    ("兒少性剝削",         "Child and Youth Sexual Exploitation"),
    ("洗錢防制法",         "Money Laundering Control Act"),
    ("其他",               "Other"),
]

# ── 警方提供資訊（上方 checkboxes）─────────────────────────────────────────
PROVIDE_ITEMS: list[tuple[str, str]] = [
    ("用戶名稱",  "Username"),
    ("電子郵件地址", "Email Address"),
    ("電話號碼",  "Phone Number"),
    ("英文姓名",  "First or Last name (other identifiers need to be included)"),
    ("錢包位址",  "Cryptocurrency Address (the transaction hash is needed to use as a confirming datapoint)"),
]

# ── 要求交易所提供資訊（下方 checkboxes）────────────────────────────────────
REQUEST_ITEMS: list[tuple[str, str]] = [
    ("電話號碼",   "Phone Number"),
    ("英文姓名",   "First or Last name (other identifiers need to be included)"),
    ("交易哈希",   "Transaction hashes"),
    ("其他",       "Additional Information"),
]

ATTACHMENT_OPTIONS: list[str] = [
    "被害人/檢舉人筆錄",
    "偵查報告",
    "幣流分析圖或報告",
    "其他",
]

# ── 法條全文（英文）──────────────────────────────────────────────────────────
_ARTICLE_228 = [
    "If a public prosecutor, because of complaint, report, voluntary surrender, "
    "or other reason, knows there is a suspicion of an offense having been committed, "
    "he shall immediately begin an investigation.",

    "In conducting the investigation referred to in the preceding section a public "
    "prosecutor may set up a period of time and order the public prosecuting affairs "
    "official, judicial police officer specified in Article 230, or judicial policeman "
    "specified in Article 231 to investigate the circumstances of the offense, to collect "
    "evidence and to submit report thereof; the case file and evidence may be delivered "
    "thereto at the same time if necessary.",

    "In the course of an investigation, an accused shall not be first summoned or "
    "interrogated unless necessary.",

    "An accused who appears by complying with a summons, voluntary surrender, or on his "
    "free will may be released on bail, or to the custody of another, or with a limitation "
    "on his residence, if the public prosecutor, after examining the accused, considers that "
    "one of the circumstances specified in the items of section I of Article 101 or the "
    "items of section I of Article 101-1 exists but application for detention is unnecessary, "
    "provided that if detention is considered necessary, the accused may be arrested without "
    "a warrant, and be informed of the fact thereof followed by an application for detention "
    "filed with the court. The provisions of sections II, III and V of Article 93 shall "
    "apply mutatis mutandis to this section.",

    "Note: Articles 1 through 343 were amended lastly on February 6, 2003.",
]

_ARTICLE_229 = [
    "Each of the following officials shall act as judicial police officer within his "
    "respective judicial district and has the duty and power of assisting a public "
    "prosecutor in investigating an offense:",

    "(1) Director General of National Police Agency, Commissioner of Police Department, "
    "General Commander of Peace Preservation Police Corps;",

    "(2) A military police superior;",

    "(3) A person authorized by law to exercise the duty and power of a judicial police "
    "officer, as specified in the preceding two items, in special matters.",

    "The judicial police officer specified in the preceding section shall send the result "
    "of the investigation to the public prosecutor; if the said officer has taken the "
    "custody of the suspect arrested with or without a warrant, he shall send the suspect "
    "to the competent public prosecutor unless otherwise provided by the law, provided that "
    "if ordered by the public prosecutor, the suspect shall be sent immediately.",

    "An accused, or suspect shall not be sent without first being arrested with or without "
    "a warrant.",

    "Note: Articles 1 through 343 were amended lastly on February 6, 2003.",
]

_ARTICLE_230 = [
    "Each of the following officials is considered to be a judicial police officer and "
    "shall obey the instructions of a public prosecutor in investigating an offense:",

    "(1) A commissioned police officer;",

    "(2) A military police officer or petty officer;",

    "(3) A person authorized by law to exercise the duty and power of a judicial police "
    "officer in special matters.",

    "The judicial police officer specified in the preceding section who suspects that an "
    "offense has been committed shall initiate an investigation immediately and report the "
    "results thereof to the competent public prosecutor and the judicial police officer "
    "referred to in the preceding article.",

    "The scene of the crime may be closed to public and inspection taken immediately, if "
    "it is necessary for investigation specified in the preceding section.",

    "Note: Articles 1 through 343 were amended lastly on February 6, 2003.",
]

_ARTICLE_231 = [
    "Each of the following officials is considered to be a judicial policeman and shall "
    "obey the orders of a public prosecutor or judicial police officer in investigating "
    "an offense:",

    "(1) A policeman;",

    "(2) A military policeman;",

    "(3) A person authorized by law to exercise the duty and power of a judicial policeman "
    "in special matters.",

    "A judicial policeman who suspects that an offense has been committed shall initiate "
    "an investigation immediately and report the results thereof to the competent public "
    "prosecutor and judicial police officer.",

    "The scene of the crime may be closed to the public and inspection taken immediately, "
    "if it is necessary for investigation specified in the preceding section.",

    "Note: Articles 1 through 343 were amended lastly on February 6, 2003.",
]


def _today_str() -> str:
    return datetime.date.today().strftime("%Y-%m-%d")


def _segment_text(text: str) -> list[tuple[str, str]]:
    """
    將文字切分為 (type, segment) 串列。
    type: 'box'   → ■□ → 新細明體
          'cjk'   → 繁體中文 + 中文標點 → 標楷體
          'latin' → 英數字母 → Consolas
    """
    def _t(ch: str) -> str:
        if ch in ('■', '□'):
            return 'box'
        cp = ord(ch)
        if (0x4E00 <= cp <= 0x9FFF or   # CJK Unified Ideographs
                0x3400 <= cp <= 0x4DBF or   # CJK Extension A
                0xF900 <= cp <= 0xFAFF or   # CJK Compatibility
                0x3000 <= cp <= 0x303F or   # CJK Symbols & Punctuation（。、…）
                0xFF00 <= cp <= 0xFFEF or   # Fullwidth forms（！？（）：；，）
                0xFE30 <= cp <= 0xFE4F or   # CJK Compatibility Forms
                0x2E80 <= cp <= 0x2EFF):    # CJK Radicals Supplement
            return 'cjk'
        return 'latin'

    if not text:
        return []
    segs: list[tuple[str, str]] = []
    cur_t = _t(text[0])
    cur_s = text[0]
    for ch in text[1:]:
        nt = _t(ch)
        if nt == cur_t:
            cur_s += ch
        else:
            segs.append((cur_t, cur_s))
            cur_t, cur_s = nt, ch
    segs.append((cur_t, cur_s))
    return segs


# ═════════════════════════════════════════════════════════════════════════════
# DOCX 生成器（精確依照範本格式）
# ═════════════════════════════════════════════════════════════════════════════

def build_docx(data: dict, out_path: str) -> None:
    """產製 Word（.docx）格式調閱申請單，格式完全依照 OKX 官方範本"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── 頁面設定 ──
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

    # 移除預設段落間距
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(4)

    def _add_runs(para, text: str, size: float, bold: bool):
        """逐段加入 Run，依字元類型套用不同字型（標楷體/Consolas/新細明體）"""
        for seg_type, seg_text in _segment_text(text):
            r = para.add_run(seg_text)
            r.bold = bold
            r.font.size = Pt(size)
            if seg_type == 'box':
                r.font.name = '新細明體'
                rPr = r._element.get_or_add_rPr()
                rPr.get_or_add_rFonts().set(qn('w:eastAsia'), '新細明體')
            elif seg_type == 'cjk':
                r.font.name = '標楷體'
                rPr = r._element.get_or_add_rPr()
                rPr.get_or_add_rFonts().set(qn('w:eastAsia'), '標楷體')
            else:
                r.font.name = 'Consolas'
                rPr = r._element.get_or_add_rPr()
                rPr.get_or_add_rFonts().set(qn('w:eastAsia'), 'Consolas')

    def _p(text: str = "", bold: bool = False, size: float = 11,
           align=WD_ALIGN_PARAGRAPH.LEFT, sb: float = 0, sa: float = 4,
           style: str = "Normal") -> object:
        para = doc.add_paragraph(style=style)
        para.alignment = align
        para.paragraph_format.space_before = Pt(sb)
        para.paragraph_format.space_after  = Pt(sa)
        if text:
            _add_runs(para, text, size, bold)
        return para

    def _blank(n: int = 1):
        for _ in range(n):
            _p(sa=2)

    # ────────────────────────────────────────────────
    # 1. Logo + 標頭：機關 + 地址（中英文）
    # ────────────────────────────────────────────────
    logo_path = data.get("logo_path") or _find_default_logo()
    if logo_path and os.path.isfile(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.paragraph_format.space_before = Pt(0)
        logo_para.paragraph_format.space_after  = Pt(4)
        logo_para.add_run().add_picture(logo_path, width=Cm(2.5))

    agency    = data.get("sender_agency",     "")
    agency_en = data.get("sender_agency_en",  "")
    addr      = data.get("sender_address",    "")
    addr_en   = data.get("sender_address_en", "")

    if agency:
        _p(agency,    bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    if addr:
        _p(addr,      size=11,            align=WD_ALIGN_PARAGRAPH.CENTER)
    if agency_en:
        _p(agency_en, size=11,            align=WD_ALIGN_PARAGRAPH.CENTER)
    if addr_en:
        _p(addr_en,   size=11,            align=WD_ALIGN_PARAGRAPH.CENTER)

    _blank()

    # ────────────────────────────────────────────────
    # 2. 標題（置中）
    # ────────────────────────────────────────────────
    _p("加密貨幣交易所調閱案件申請單",
       bold=True, size=14,
       align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=8)

    # ────────────────────────────────────────────────
    # 3. 受文者資訊
    # ────────────────────────────────────────────────
    recipient = data.get("recipient_name",  "")
    rec_email = data.get("recipient_email", "")
    doc_date  = data.get("doc_date",   _today_str())
    doc_num   = data.get("doc_number", "")

    _p(f"受文者：{recipient}")
    _p(f"To: {recipient}")
    _p(f"Date:{doc_date}")
    if doc_num:
        _p(f"Official Ref. No.: {doc_num}")
    if rec_email:
        _p(f"Via email to: {rec_email}")

    _blank()

    # ────────────────────────────────────────────────
    # 4. 案件性質（勾選格式，依範本排列）
    # ────────────────────────────────────────────────
    _p("因偵辦下列案件（請勾選所偵辦案件之性質，可複選）"
       "The Criminal Code of Taiwan which has been violated：", sa=2)

    selected_types = set(data.get("case_types", []))

    def _mark(cn: str) -> str:
        return "■" if cn in selected_types else "□"

    # 前6項三個一排（2列各3項）
    _three = [
        (CASE_TYPES[0], CASE_TYPES[1], CASE_TYPES[2]),
        (CASE_TYPES[3], CASE_TYPES[4], CASE_TYPES[5]),
    ]
    for row in _three:
        parts = []
        for cn, en in row:
            cell_text = f"{_mark(cn)}{cn}{en}"
            parts.append(f"{cell_text:<28}")
        _p("  ".join(parts), sa=1)

    # 第7-8項兩個一排
    cn7, en7 = CASE_TYPES[6]
    cn8, en8 = CASE_TYPES[7]
    _p(f"{_mark(cn7)}{cn7}{en7}      {_mark(cn8)}{cn8}{en8}", sa=1)

    # 第9-14項每行一個
    for cn, en in CASE_TYPES[8:]:
        if cn == "兒少性剝削":
            _p(f"{_mark(cn)}{cn}{en} (兒少類須註明被害人年齡)", sa=1)
        elif cn == "其他":
            _p(f"{_mark(cn)}{cn}other (請填寫 please specify)：", sa=1)
        else:
            _p(f"{_mark(cn)}{cn}{en}", sa=1)

    _blank()

    # ────────────────────────────────────────────────
    # 5. 案情簡述
    # ────────────────────────────────────────────────
    _p("1、案情簡述(請以中文及英文簡述人事時地物及與業者之關聯) "
       "Brief description of investigation：", sa=2)

    desc_cn = data.get("desc_cn", "")
    desc_en = data.get("desc_en", "")
    if desc_cn:
        _p(f"中文: {desc_cn}", sa=4)
    if desc_en:
        _p(f"英文: {desc_en}", sa=4)

    _blank()

    # ────────────────────────────────────────────────
    # 6. 法律依據（中文條文 + 英文全文）
    # ────────────────────────────────────────────────
    _p("2、依據法條：刑事訴訟法第228、第229條、第230條及第231條", sa=2)
    _p("Article 228, 229, 230 and 231 of Code of Criminal Procedure", sa=4)

    for art_num, art_paras in [
        ("Article 228", _ARTICLE_228),
        ("Article 229", _ARTICLE_229),
        ("Article 230", _ARTICLE_230),
        ("Article 231", _ARTICLE_231),
    ]:
        _p(art_num, bold=True, sa=2)
        for para_text in art_paras:
            _p(para_text, size=10, sa=2)
        _blank()

    _p(f"We understand that {recipient} is not within our jurisdiction, "
       f"but we will appreciate any voluntary assistance that {recipient} "
       f"may be able to provide.", sa=4)

    # ────────────────────────────────────────────────
    # 7. 警方提供資訊 + 錢包地址
    # ────────────────────────────────────────────────
    _p("3、警方提供資訊（Provided Information and action）：", sa=2)

    provided = set(data.get("provided_items", ["錢包位址"]))
    for cn, en in PROVIDE_ITEMS:
        mk = "■" if cn in provided else "□"
        if cn == "錢包位址":
            _p(f"{mk}{cn}（註明幣別及交易雜湊值）{en}", sa=1)
        else:
            _p(f"{mk}{cn}{en}", sa=1)

    wallets: list[dict] = data.get("wallets", [])
    if wallets:
        _blank()
        _p("1.Please provide KYC of every account interacted with the suspicious "
           "addresses below 請提供與以下可疑地址有往來的所有帳戶的KYC 認證信息:", sa=2)

        # 依幣種/鏈分組
        seen_chains: list[str] = []
        for wlt in wallets:
            chain = wlt.get("chain", "")
            if chain not in seen_chains:
                seen_chains.append(chain)

        for chain in seen_chains:
            chain_wallets = [w for w in wallets if w.get("chain") == chain]
            _p(f"For {chain}:", sa=1)
            for wlt in chain_wallets:
                addr_line = wlt.get("address", "")
                tx_hash   = wlt.get("tx_hash", "").strip()
                if tx_hash:
                    addr_line += f"\n(tx hash: {tx_hash})"
                _p(addr_line, size=10, sa=2)

    _blank()

    requested = set(data.get("requested_items", ["電話號碼", "英文姓名", "交易哈希"]))
    for cn, en in REQUEST_ITEMS:
        mk = "■" if cn in requested else "□"
        if cn == "英文姓名":
            _p(f"{mk}英文姓名First or Last name (other identifiers need to be included)",
               sa=1)
        elif cn == "其他":
            _p(f"{mk}其他Additional Information＿＿＿＿＿＿＿＿＿＿＿", sa=1)
        else:
            _p(f"{mk}{cn}{en}", sa=1)

    _blank()

    # ────────────────────────────────────────────────
    # 8. 調閱時間區間
    # ────────────────────────────────────────────────
    _p("4、調閱時間區間(Duration of request)：", sa=2)
    date_from = data.get("date_from", "")
    date_to   = data.get("date_to",   "")
    _p(f"始期(From)：{date_from}", sa=1)
    _p(f"終期(To)  ：{date_to}",   sa=4)

    # ────────────────────────────────────────────────
    # 9. 附件
    # ────────────────────────────────────────────────
    _p("5、檢附佐證及其他相關資料Evidence or other relevant Information：", sa=2)
    attachments = data.get("attachments", [])
    att_map = {
        "被害人/檢舉人筆錄": "□被害人/檢舉人筆錄Victim or witness's  statement",
        "偵查報告":          "□偵查報告Report of investigation",
        "幣流分析圖或報告":  "□幣流分析圖或報告 Tracing report",
        "其他":              "□其他：＿＿＿",
    }
    for att_cn, att_line in att_map.items():
        mark = "■" if att_cn in attachments else "□"
        _p(mark + att_line[1:], sa=1)

    _blank()

    # ────────────────────────────────────────────────
    # 10. 不披露請求
    # ────────────────────────────────────────────────
    nd_date = data.get("nondisclosure_date", "")
    nd_text = ("6、暫不向調閱對象披露資訊至Non-disclosure requests："
               "Please do not to disclose information to the user until")
    if nd_date:
        nd_text += f" {nd_date}"
    _p(nd_text, sa=4)

    # ────────────────────────────────────────────────
    # 11. 特殊請求
    # ────────────────────────────────────────────────
    keep_open     = data.get("keep_account_open", False)
    special_notes = data.get("special_notes", "")

    _p(f"7、特殊請求（請注明。特殊請求的自願協助需經 {recipient} 考慮，並可能受到進一步限制，"
       f"具體條件將根據個案決定。該條件將是{recipient} 自願執行的前提。）"
       f"Special request (Please specify. The voluntary assistance of a special request "
       f"is subject to {recipient}'s consideration and conditions will be proposed on a "
       f"case-by-case basis, acceptance of which is a condition to its voluntary "
       f"implementation.):", sa=2)

    mk_keep = "■" if keep_open else "□"
    _p(f"{mk_keep}保持開啟的要求(即使發現非法活動)"
       f"Please keep the account open (even if illegal activity is discovered)", sa=2)

    _p(f"(是否需要該帳號保持開啟，否則交易所可能將其凍結, 如不需要，請刪除此項。"
       f"請注意，{recipient} 只能將帳號保持開啟7天。{recipient}保留根據{recipient}服務條款"
       f"凍結或終止帳號的權利，包括在沒有執法機構要求的情況下凍結或終止帳戶，"
       f"例如因合規原因或為了保護 {recipient} 用戶和平台的安全而進行的凍結。"
       f"[Please specify if the account is required to be kept open. Otherwise, please "
       f"remove this part. Please note that {recipient} can only keep the account open "
       f"for 7 days. {recipient} reserves the right to freeze or terminate the said "
       f"account in accordance with {recipient} Terms of Service, including to freeze or "
       f"terminate an account independent of a law enforcement request, such as freeze "
       f"imposed due to compliance reason or to protect the safety of {recipient} users "
       f"and platform.]", size=10, sa=4)

    if special_notes:
        _p(special_notes, sa=4)

    _p("8、如有可能請提供電子檔案如CSV或EXCEL等可編輯文件。"
       "Please provide editable document as possible as it could be, such as csv or excel.",
       sa=8)

    # ────────────────────────────────────────────────
    # 12. 聯絡資訊
    # ────────────────────────────────────────────────
    _p("9、聯絡資訊：", bold=True, sa=2)
    _p("機關Department：", sa=1)

    name     = data.get("sender_name",      "")
    name_en  = data.get("sender_name_en",   "")
    unit     = data.get("sender_unit",      "")
    title    = data.get("sender_title",     "")
    phone    = data.get("sender_phone",     "")
    email    = data.get("sender_email",     "")

    agency_line = agency
    if agency_en:
        agency_line += agency_en
    _p(agency_line, sa=1)
    if unit:
        _p(f"單位Unit：{unit}", sa=1)
    if title:
        _p(f"職稱Title：{title}", sa=1)
    name_line = f"姓名Name：{name}"
    if name_en:
        name_line += f"  {name_en}"
    _p(name_line, sa=1)
    if phone:
        _p(f"電話Phone：{phone}", sa=1)
    if email:
        _p(f"電郵Email：{email}", sa=1)

    _blank()
    _p("請查照。", sa=1)
    _p("Thanks.", sa=4)
    _p("Name, title and chop of authorized official:", sa=8)

    _p("我等同意把聯絡資訊提供給被凍結帳號的持有人（如果有要求暫不披露，"
       "則只同意在不披露請求失效後披露）。"
       "We agree to sharing our contact details with the holder of the frozen account(s) "
       "(if non-disclosure is requested, then we only agree to sharing our contact details "
       "after the non-disclosure request has lapsed).", size=10)

    doc.save(out_path)


# ═════════════════════════════════════════════════════════════════════════════
# ODT 生成器（與 DOCX 相同結構）
# ═════════════════════════════════════════════════════════════════════════════

def build_odt(data: dict, out_path: str) -> None:
    """產製 OpenDocument Text（.odt）格式調閱申請單"""
    from odf.opendocument import OpenDocumentText
    from odf.style import Style, TextProperties, ParagraphProperties, TableColumnProperties
    from odf.text import P, H, Span
    from odf.table import Table, TableRow, TableCell, TableColumn
    from odf.namespaces import FONS, STYLENS

    doc = OpenDocumentText()

    def _make_para_style(name, size="11pt", bold=False, align="left",
                         color=None, margintop="0cm", marginbottom="0.15cm"):
        s = Style(name=name, family="paragraph")
        pp: dict[str, Any] = {"textalign": align, "margintop": margintop,
                               "marginbottom": marginbottom}
        s.addElement(ParagraphProperties(**pp))
        tp_kw: dict[str, Any] = {"fontsize": size, "fontfamily": "標楷體",
                                  "fontnamecomplex": "標楷體", "fontfamilycomplex": "標楷體"}
        if bold:  tp_kw["fontweight"] = "bold"
        if color: tp_kw["color"]      = color
        tp_el = TextProperties(**tp_kw)
        tp_el.attributes[(STYLENS, 'font-name-asian')]   = '標楷體'
        tp_el.attributes[(FONS,    'font-family-asian')] = '標楷體'
        s.addElement(tp_el)
        doc.styles.addElement(s)

    _make_para_style("Normal",    size="11pt")
    _make_para_style("Bold",      size="11pt", bold=True)
    _make_para_style("Title",     size="14pt", bold=True, align="center",
                     margintop="0.3cm", marginbottom="0.4cm")
    _make_para_style("Small",     size="10pt")
    _make_para_style("Mono",      size="10pt")

    # Inline text styles for per-character font switching
    def _make_text_style(name: str, font: str) -> None:
        s = Style(name=name, family="text")
        tp = TextProperties(fontname=font, fontfamily=font,
                            fontnamecomplex=font, fontfamilycomplex=font)
        tp.attributes[(STYLENS, 'font-name-asian')]   = font
        tp.attributes[(FONS,    'font-family-asian')] = font
        s.addElement(tp)
        doc.automaticstyles.addElement(s)

    _make_text_style("T_CJK",   "標楷體")
    _make_text_style("T_Latin", "Consolas")
    _make_text_style("T_Box",   "新細明體")

    _col_counter = [0]
    def _add_col_style(width: str) -> str:
        _col_counter[0] += 1
        name = f"ColW{_col_counter[0]}"
        s = Style(name=name, family="table-column")
        s.addElement(TableColumnProperties(columnwidth=width))
        doc.automaticstyles.addElement(s)
        return name

    _ODT_STYLE = {"cjk": "T_CJK", "latin": "T_Latin", "box": "T_Box"}

    def _add_p(text: str, style: str = "Normal") -> None:
        p = P(stylename=style)
        for seg_type, seg_text in _segment_text(text):
            span = Span(stylename=_ODT_STYLE[seg_type])
            span.addText(seg_text)
            p.addElement(span)
        doc.text.addElement(p)

    def _blank():
        _add_p("", "Normal")

    # 1. Logo + 標頭
    logo_path = data.get("logo_path") or _find_default_logo()
    if logo_path and os.path.isfile(logo_path):
        from odf.draw import Frame, Image as OdfImage
        from odf.style import Style as OdfStyle, GraphicProperties

        # 圖形樣式物件（必須傳入物件，不能傳字串）
        gfx_style = OdfStyle(name="LogoFrame", family="graphic")
        gfx_style.addElement(GraphicProperties(stroke="none", fill="none"))
        doc.automaticstyles.addElement(gfx_style)

        logo_href = doc.addPicture(logo_path)
        frame = Frame(stylename=gfx_style,
                      width="2.5cm", height="2.5cm",
                      anchortype="as-char")
        frame.addElement(OdfImage(href=logo_href, type="simple"))
        logo_p = P(stylename="Title")
        logo_p.addElement(frame)
        doc.text.addElement(logo_p)

    agency    = data.get("sender_agency",     "")
    agency_en = data.get("sender_agency_en",  "")
    addr      = data.get("sender_address",    "")
    addr_en   = data.get("sender_address_en", "")
    if agency:    _add_p(agency,    "Title")
    if addr:      _add_p(addr,      "Normal")
    if agency_en: _add_p(agency_en, "Normal")
    if addr_en:   _add_p(addr_en,   "Normal")
    _blank()

    # 2. 標題
    _add_p("加密貨幣交易所調閱案件申請單", "Title")

    # 3. 受文者
    recipient = data.get("recipient_name",  "")
    rec_email = data.get("recipient_email", "")
    doc_date  = data.get("doc_date",   _today_str())
    doc_num   = data.get("doc_number", "")
    _add_p(f"受文者：{recipient}")
    _add_p(f"To: {recipient}")
    _add_p(f"Date:{doc_date}")
    if doc_num:   _add_p(f"Official Ref. No.: {doc_num}")
    if rec_email: _add_p(f"Via email to: {rec_email}")
    _blank()

    # 4. 案件性質
    _add_p("因偵辦下列案件（請勾選所偵辦案件之性質，可複選）"
           "The Criminal Code of Taiwan which has been violated：")
    selected_types = set(data.get("case_types", []))

    def _mark(cn: str) -> str:
        return "■" if cn in selected_types else "□"

    for row in [(CASE_TYPES[0], CASE_TYPES[1], CASE_TYPES[2]),
                (CASE_TYPES[3], CASE_TYPES[4], CASE_TYPES[5])]:
        _add_p("  ".join(f"{_mark(cn)}{cn}{en}" for cn, en in row))
    cn7, en7 = CASE_TYPES[6]; cn8, en8 = CASE_TYPES[7]
    _add_p(f"{_mark(cn7)}{cn7}{en7}      {_mark(cn8)}{cn8}{en8}")
    for cn, en in CASE_TYPES[8:]:
        if cn == "兒少性剝削":
            _add_p(f"{_mark(cn)}{cn}{en} (兒少類須註明被害人年齡)")
        elif cn == "其他":
            _add_p(f"{_mark(cn)}{cn}other (請填寫 please specify)：")
        else:
            _add_p(f"{_mark(cn)}{cn}{en}")
    _blank()

    # 5. 案情簡述
    _add_p("1、案情簡述(請以中文及英文簡述人事時地物及與業者之關聯) "
           "Brief description of investigation：")
    desc_cn = data.get("desc_cn", "")
    desc_en = data.get("desc_en", "")
    if desc_cn: _add_p(f"中文: {desc_cn}")
    if desc_en: _add_p(f"英文: {desc_en}")
    _blank()

    # 6. 法律依據
    _add_p("2、依據法條：刑事訴訟法第228、第229條、第230條及第231條")
    _add_p("Article 228, 229, 230 and 231 of Code of Criminal Procedure")
    for art_num, art_paras in [("Article 228", _ARTICLE_228), ("Article 229", _ARTICLE_229),
                                ("Article 230", _ARTICLE_230), ("Article 231", _ARTICLE_231)]:
        _add_p(art_num, "Bold")
        for pt in art_paras:
            _add_p(pt, "Small")
        _blank()
    _add_p(f"We understand that {recipient} is not within our jurisdiction, "
           f"but we will appreciate any voluntary assistance that {recipient} "
           f"may be able to provide.")
    _blank()

    # 7. 警方提供資訊
    _add_p("3、警方提供資訊（Provided Information and action）：")
    provided = set(data.get("provided_items", ["錢包位址"]))
    for cn, en in PROVIDE_ITEMS:
        mk = "■" if cn in provided else "□"
        if cn == "錢包位址":
            _add_p(f"{mk}{cn}（註明幣別及交易雜湊值）{en}")
        else:
            _add_p(f"{mk}{cn}{en}")

    wallets: list[dict] = data.get("wallets", [])
    if wallets:
        _blank()
        _add_p("1.Please provide KYC of every account interacted with the suspicious "
               "addresses below 請提供與以下可疑地址有往來的所有帳戶的KYC 認證信息:")
        seen_chains: list[str] = []
        for wlt in wallets:
            if wlt.get("chain") not in seen_chains:
                seen_chains.append(wlt["chain"])
        for chain in seen_chains:
            _add_p(f"For {chain}:", "Bold")
            for wlt in wallets:
                if wlt.get("chain") == chain:
                    _add_p(wlt.get("address", ""), "Mono")
                    if wlt.get("tx_hash"):
                        _add_p(f"(tx hash: {wlt['tx_hash']})", "Small")
    _blank()

    requested = set(data.get("requested_items", ["電話號碼", "英文姓名", "交易哈希"]))
    for cn, en in REQUEST_ITEMS:
        mk = "■" if cn in requested else "□"
        if cn == "英文姓名":
            _add_p(f"{mk}英文姓名First or Last name (other identifiers need to be included)")
        elif cn == "其他":
            _add_p(f"{mk}其他Additional Information＿＿＿＿＿＿＿＿＿＿＿")
        else:
            _add_p(f"{mk}{cn}{en}")
    _blank()

    # 8. 調閱時間
    _add_p("4、調閱時間區間(Duration of request)：")
    _add_p(f"始期(From)：{data.get('date_from', '')}")
    _add_p(f"終期(To)  ：{data.get('date_to',   '')}")
    _blank()

    # 9. 附件
    _add_p("5、檢附佐證及其他相關資料Evidence or other relevant Information：")
    attachments = data.get("attachments", [])
    att_map = [
        ("被害人/檢舉人筆錄", "被害人/檢舉人筆錄Victim or witness's  statement"),
        ("偵查報告",          "偵查報告Report of investigation"),
        ("幣流分析圖或報告",  "幣流分析圖或報告 Tracing report"),
        ("其他",              "其他：＿＿＿"),
    ]
    for att_cn, att_full in att_map:
        mk = "■" if att_cn in attachments else "□"
        _add_p(f"{mk}{att_full}")
    _blank()

    # 10. 不披露
    nd_date = data.get("nondisclosure_date", "")
    nd_text = ("6、暫不向調閱對象披露資訊至Non-disclosure requests："
               "Please do not to disclose information to the user until")
    if nd_date:
        nd_text += f" {nd_date}"
    _add_p(nd_text)
    _blank()

    # 11. 特殊請求
    keep_open     = data.get("keep_account_open", False)
    special_notes = data.get("special_notes", "")
    _add_p(f"7、特殊請求（請注明。特殊請求的自願協助需經 {recipient} 考慮..."
           f"）Special request (...):")
    mk_keep = "■" if keep_open else "□"
    _add_p(f"{mk_keep}保持開啟的要求(即使發現非法活動)"
           f"Please keep the account open (even if illegal activity is discovered)")
    _add_p(f"(請注意，{recipient} 只能將帳號保持開啟7天。)", "Small")
    if special_notes:
        _add_p(special_notes)
    _add_p("8、如有可能請提供電子檔案如CSV或EXCEL等可編輯文件。"
           "Please provide editable document as possible as it could be, such as csv or excel.")
    _blank()

    # 12. 聯絡資訊
    _add_p("9、聯絡資訊：", "Bold")
    _add_p("機關Department：")
    agency_line = agency + (agency_en if agency_en else "")
    _add_p(agency_line)
    name    = data.get("sender_name",    "")
    name_en = data.get("sender_name_en", "")
    unit    = data.get("sender_unit",    "")
    title   = data.get("sender_title",   "")
    phone   = data.get("sender_phone",   "")
    email   = data.get("sender_email",   "")
    if unit:  _add_p(f"單位Unit：{unit}")
    if title: _add_p(f"職稱Title：{title}")
    name_line = f"姓名Name：{name}"
    if name_en: name_line += f"  {name_en}"
    _add_p(name_line)
    if phone: _add_p(f"電話Phone：{phone}")
    if email: _add_p(f"電郵Email：{email}")
    _blank()
    _add_p("請查照。")
    _add_p("Thanks.")
    _blank()
    _add_p("Name, title and chop of authorized official:")
    _blank()
    _add_p("我等同意把聯絡資訊提供給被凍結帳號的持有人（如果有要求暫不披露，"
           "則只同意在不披露請求失效後披露）。"
           "We agree to sharing our contact details with the holder of the frozen account(s) "
           "(if non-disclosure is requested, then we only agree to sharing our contact details "
           "after the non-disclosure request has lapsed).", "Small")

    doc.save(out_path)


# ═════════════════════════════════════════════════════════════════════════════
# PDF 生成器（reportlab，與 DOCX 相同結構）
# ═════════════════════════════════════════════════════════════════════════════

def _find_cjk_font() -> tuple[str, int] | None:
    candidates = [
        (r"C:\Windows\Fonts\msjh.ttc",   0),
        (r"C:\Windows\Fonts\mingliu.ttc", 0),
        (r"C:\Windows\Fonts\simsun.ttc",  0),
        (r"C:\Windows\Fonts\kaiu.ttf",    0),
    ]
    for path, idx in candidates:
        if os.path.isfile(path):
            return path, idx
    return None


def build_pdf(data: dict, out_path: str) -> None:
    """產製 PDF 格式調閱申請單（reportlab）"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, HRFlowable)
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise ImportError("請先安裝 reportlab：pip install reportlab")

    # ── 字型載入：標楷體 / Consolas / 新細明體 ──
    _PDF_FONTS: dict[str, str] = {}

    def _try_register(alias: str, path: str, subfont: int = 0) -> bool:
        if not os.path.isfile(path):
            return False
        try:
            pdfmetrics.registerFont(TTFont(alias, path, subfontIndex=subfont))
            return True
        except Exception:
            return False

    # 標楷體（主要 CJK 字型）
    if _try_register("KaiuFont", r"C:\Windows\Fonts\kaiu.ttf"):
        _PDF_FONTS["cjk"] = "KaiuFont"
        _try_register("KaiuFont-Bold", r"C:\Windows\Fonts\kaiu.ttf")  # 相同檔案作為 Bold 別名
    else:
        fi = _find_cjk_font()
        if fi is None:
            raise RuntimeError("找不到 Windows 中文字型\n請改用 DOCX 或 ODT 格式。")
        _try_register("KaiuFont", fi[0], fi[1])
        _try_register("KaiuFont-Bold", fi[0], fi[1])
        _PDF_FONTS["cjk"] = "KaiuFont"

    # Consolas（Latin 字型）
    if _try_register("ConsolasFont", r"C:\Windows\Fonts\consola.ttf"):
        _PDF_FONTS["latin"] = "ConsolasFont"
        _try_register("ConsolasFont-Bold", r"C:\Windows\Fonts\consolab.ttf")
    else:
        _PDF_FONTS["latin"] = _PDF_FONTS["cjk"]

    # 新細明體（■□ 字型）— mingliu.ttc subfont 1
    if _try_register("MingLiuFont", r"C:\Windows\Fonts\mingliu.ttc", subfont=1):
        _PDF_FONTS["box"] = "MingLiuFont"
        _try_register("MingLiuFont-Bold", r"C:\Windows\Fonts\mingliu.ttc", subfont=1)
    else:
        _PDF_FONTS["box"] = _PDF_FONTS["cjk"]

    try:
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        registerFontFamily("KaiuFont",
                           normal="KaiuFont", bold="KaiuFont-Bold",
                           italic="KaiuFont", boldItalic="KaiuFont-Bold")
        if _PDF_FONTS["latin"] == "ConsolasFont":
            bold_alias = "ConsolasFont-Bold" if _try_register(
                "ConsolasFont-Bold-chk", r"C:\Windows\Fonts\consolab.ttf") else "ConsolasFont"
            registerFontFamily("ConsolasFont",
                               normal="ConsolasFont", bold=bold_alias,
                               italic="ConsolasFont", boldItalic=bold_alias)
    except Exception:
        pass

    _fn   = _PDF_FONTS["cjk"]

    def _xml_esc(s: str) -> str:
        return s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def _markup(text: str, bold_text: bool = False) -> str:
        """將文字轉為 reportlab XML markup，依字元類型切換字型"""
        parts = []
        for seg_type, seg_text in _segment_text(text):
            base_fn = _PDF_FONTS.get(seg_type, _fn)
            fn = base_fn + ("-Bold" if bold_text else "")
            parts.append(f'<font name="{fn}">{_xml_esc(seg_text)}</font>')
        return ''.join(parts)

    base  = ParagraphStyle("base",  fontName=_fn, fontSize=11, leading=18, spaceAfter=4)
    bold  = ParagraphStyle("bold",  fontName=_fn, fontSize=11, leading=18, spaceAfter=4,
                           textColor=colors.black)
    title = ParagraphStyle("title", fontName=_fn, fontSize=14, leading=22, spaceAfter=8,
                           alignment=1)  # center
    small = ParagraphStyle("small", fontName=_fn, fontSize=9,  leading=14, spaceAfter=3)
    mono  = ParagraphStyle("mono",  fontName=_fn, fontSize=9,  leading=14, spaceAfter=2)

    def _p(text: str, st=None) -> Paragraph:
        return Paragraph(_markup(text), st or base)

    def _b(text: str) -> Paragraph:
        return Paragraph(_markup(text, bold_text=True), bold)

    def _sp(n=6) -> Spacer:
        return Spacer(1, n)

    story = []

    # 1. Logo + 標頭
    logo_path = data.get("logo_path") or _find_default_logo()
    if logo_path and os.path.isfile(logo_path):
        from reportlab.platypus import Image as RLImage
        logo_img = RLImage(logo_path, width=2.5*cm, height=2.5*cm)
        logo_img.hAlign = "CENTER"
        story.append(logo_img)
        story.append(_sp(4))

    agency    = data.get("sender_agency",     "")
    agency_en = data.get("sender_agency_en",  "")
    addr      = data.get("sender_address",    "")
    addr_en   = data.get("sender_address_en", "")

    center = ParagraphStyle("center_hdr", fontName=_fn, fontSize=12,
                            leading=18, spaceAfter=2, alignment=1)
    if agency:    story.append(Paragraph(_markup(agency, bold_text=True), center))
    if addr:      story.append(Paragraph(_markup(addr),      center))
    if agency_en: story.append(Paragraph(_markup(agency_en), center))
    if addr_en:   story.append(Paragraph(_markup(addr_en),   center))
    story.append(_sp(8))

    # 2. 標題
    story.append(_p("加密貨幣交易所調閱案件申請單", title))

    # 3. 受文者
    recipient = data.get("recipient_name",  "")
    rec_email = data.get("recipient_email", "")
    doc_date  = data.get("doc_date",   _today_str())
    doc_num   = data.get("doc_number", "")
    story.append(_p(f"受文者：{recipient}"))
    story.append(_p(f"To: {recipient}"))
    story.append(_p(f"Date:{doc_date}"))
    if doc_num:   story.append(_p(f"Official Ref. No.: {doc_num}"))
    if rec_email: story.append(_p(f"Via email to: {rec_email}"))
    story.append(_sp(8))

    # 4. 案件性質
    story.append(_p("因偵辦下列案件（請勾選所偵辦案件之性質，可複選）"
                    "The Criminal Code of Taiwan which has been violated："))
    selected_types = set(data.get("case_types", []))

    def _mark(cn: str) -> str:
        return "■" if cn in selected_types else "□"

    for row in [(CASE_TYPES[0], CASE_TYPES[1], CASE_TYPES[2]),
                (CASE_TYPES[3], CASE_TYPES[4], CASE_TYPES[5])]:
        story.append(_p("    ".join(f"{_mark(cn)}{cn}{en}" for cn, en in row)))
    cn7, en7 = CASE_TYPES[6]; cn8, en8 = CASE_TYPES[7]
    story.append(_p(f"{_mark(cn7)}{cn7}{en7}      {_mark(cn8)}{cn8}{en8}"))
    for cn, en in CASE_TYPES[8:]:
        if cn == "兒少性剝削":
            story.append(_p(f"{_mark(cn)}{cn}{en} (兒少類須註明被害人年齡)"))
        elif cn == "其他":
            story.append(_p(f"{_mark(cn)}{cn}other (請填寫 please specify)："))
        else:
            story.append(_p(f"{_mark(cn)}{cn}{en}"))
    story.append(_sp(8))

    # 5. 案情
    story.append(_p("1、案情簡述(請以中文及英文簡述人事時地物及與業者之關聯) "
                    "Brief description of investigation："))
    desc_cn = data.get("desc_cn", "")
    desc_en = data.get("desc_en", "")
    if desc_cn: story.append(_p(f"中文: {desc_cn}"))
    if desc_en: story.append(_p(f"英文: {desc_en}"))
    story.append(_sp(8))

    # 6. 法律依據
    story.append(_p("2、依據法條：刑事訴訟法第228、第229條、第230條及第231條"))
    story.append(_p("Article 228, 229, 230 and 231 of Code of Criminal Procedure"))
    for art_num, art_paras in [("Article 228", _ARTICLE_228), ("Article 229", _ARTICLE_229),
                                ("Article 230", _ARTICLE_230), ("Article 231", _ARTICLE_231)]:
        story.append(_b(art_num))
        for pt in art_paras:
            story.append(_p(pt, small))
        story.append(_sp(4))
    story.append(_p(f"We understand that {recipient} is not within our jurisdiction, "
                    f"but we will appreciate any voluntary assistance that {recipient} "
                    f"may be able to provide."))
    story.append(_sp(8))

    # 7. 提供資訊
    story.append(_p("3、警方提供資訊（Provided Information and action）："))
    provided = set(data.get("provided_items", ["錢包位址"]))
    for cn, en in PROVIDE_ITEMS:
        mk = "■" if cn in provided else "□"
        if cn == "錢包位址":
            story.append(_p(f"{mk}{cn}（註明幣別及交易雜湊值）{en}"))
        else:
            story.append(_p(f"{mk}{cn}{en}"))

    wallets: list[dict] = data.get("wallets", [])
    if wallets:
        story.append(_sp(4))
        story.append(_p("1.Please provide KYC of every account interacted with the "
                        "suspicious addresses below 請提供與以下可疑地址有往來的所有帳戶的KYC 認證信息:"))
        seen_chains: list[str] = []
        for wlt in wallets:
            if wlt.get("chain") not in seen_chains:
                seen_chains.append(wlt["chain"])
        for chain in seen_chains:
            story.append(_b(f"For {chain}:"))
            for wlt in wallets:
                if wlt.get("chain") == chain:
                    story.append(_p(wlt.get("address", ""), mono))
                    if wlt.get("tx_hash"):
                        story.append(_p(f"(tx hash: {wlt['tx_hash']})", small))
    story.append(_sp(4))

    requested = set(data.get("requested_items", ["電話號碼", "英文姓名", "交易哈希"]))
    for cn, en in REQUEST_ITEMS:
        mk = "■" if cn in requested else "□"
        if cn == "英文姓名":
            story.append(_p(f"{mk}英文姓名First or Last name (other identifiers need to be included)"))
        elif cn == "其他":
            story.append(_p(f"{mk}其他Additional Information＿＿＿＿＿＿＿＿＿＿＿"))
        else:
            story.append(_p(f"{mk}{cn}{en}"))
    story.append(_sp(8))

    # 8-9. 時間 + 附件
    story.append(_p("4、調閱時間區間(Duration of request)："))
    story.append(_p(f"始期(From)：{data.get('date_from', '')}"))
    story.append(_p(f"終期(To)  ：{data.get('date_to', '')}"))
    story.append(_sp(8))

    attachments = data.get("attachments", [])
    story.append(_p("5、檢附佐證及其他相關資料Evidence or other relevant Information："))
    att_map = [("被害人/檢舉人筆錄", "被害人/檢舉人筆錄Victim or witness's  statement"),
               ("偵查報告",          "偵查報告Report of investigation"),
               ("幣流分析圖或報告",  "幣流分析圖或報告 Tracing report"),
               ("其他",              "其他：＿＿＿")]
    for att_cn, att_full in att_map:
        mk = "■" if att_cn in attachments else "□"
        story.append(_p(f"{mk}{att_full}"))
    story.append(_sp(8))

    # 10-11. 不披露 + 特殊請求
    nd_date = data.get("nondisclosure_date", "")
    nd_text = ("6、暫不向調閱對象披露資訊至Non-disclosure requests："
               "Please do not to disclose information to the user until")
    if nd_date:
        nd_text += f" {nd_date}"
    story.append(_p(nd_text))
    story.append(_sp(8))

    keep_open     = data.get("keep_account_open", False)
    special_notes = data.get("special_notes", "")
    story.append(_p(f"7、特殊請求（...）Special request (...):"))
    mk_keep = "■" if keep_open else "□"
    story.append(_p(f"{mk_keep}保持開啟的要求(即使發現非法活動)"
                    f"Please keep the account open (even if illegal activity is discovered)"))
    story.append(_p(f"(請注意，{recipient} 只能將帳號保持開啟7天。)", small))
    if special_notes:
        story.append(_p(special_notes))
    story.append(_p("8、如有可能請提供電子檔案如CSV或EXCEL等可編輯文件。"
                    "Please provide editable document as possible as it could be, such as csv or excel."))
    story.append(_sp(12))

    # 12. 聯絡資訊
    story.append(_b("9、聯絡資訊："))
    story.append(_p("機關Department："))
    agency_line = agency + (agency_en if agency_en else "")
    story.append(_p(agency_line))
    name    = data.get("sender_name",    "")
    name_en = data.get("sender_name_en", "")
    unit    = data.get("sender_unit",    "")
    title   = data.get("sender_title",   "")
    phone   = data.get("sender_phone",   "")
    email   = data.get("sender_email",   "")
    if unit:  story.append(_p(f"單位Unit：{unit}"))
    if title: story.append(_p(f"職稱Title：{title}"))
    name_line = f"姓名Name：{name}"
    if name_en: name_line += f"  {name_en}"
    story.append(_p(name_line))
    if phone: story.append(_p(f"電話Phone：{phone}"))
    if email: story.append(_p(f"電郵Email：{email}"))
    story.append(_sp(8))
    story.append(_p("請查照。"))
    story.append(_p("Thanks."))
    story.append(_sp(16))
    story.append(_p("Name, title and chop of authorized official:"))
    story.append(_sp(16))
    story.append(_p("我等同意把聯絡資訊提供給被凍結帳號的持有人…"
                    "We agree to sharing our contact details with the holder of the frozen account(s)…", small))

    SimpleDocTemplate(out_path, pagesize=A4,
                      topMargin=2.5*cm, bottomMargin=2.5*cm,
                      leftMargin=3.0*cm, rightMargin=2.5*cm).build(story)


# ═════════════════════════════════════════════════════════════════════════════
# 統一入口
# ═════════════════════════════════════════════════════════════════════════════

def build_inquiry(data: dict, out_path: str) -> None:
    ext = os.path.splitext(out_path)[1].lower()
    if ext == ".docx":
        build_docx(data, out_path)
    elif ext == ".odt":
        build_odt(data, out_path)
    elif ext == ".pdf":
        build_pdf(data, out_path)
    else:
        raise ValueError(f"不支援的格式：{ext}（請使用 .docx / .odt / .pdf）")
