"""
flow_report_builder.py
幣流分析報告（司法鑑識專家意見書）Word 文件生成器。
格式依據：幣流報告格式黃瑋琳案(稿).docx
"""
from __future__ import annotations
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# 樣式輔助（與 case_template_builder 共用邏輯）
# ─────────────────────────────────────────────────────────────────────────────

def _font(run, size=10.5, bold=False, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Microsoft JhengHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _p(doc, text='', bold=False, size=10.5, color=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4, italic=False):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(sb)
    para.paragraph_format.space_after  = Pt(sa)
    if text:
        r = para.add_run(text)
        _font(r, size=size, bold=bold, color=color, italic=italic)
    return para


def _h1(doc, text):
    _p(doc, text, bold=True, size=13, color='1F3864', sb=10, sa=4)


def _h2(doc, text):
    _p(doc, text, bold=True, size=11.5, color='2E75B6', sb=6, sa=3)


def _h3(doc, text):
    _p(doc, text, bold=True, size=10.5, color='404040', sb=4, sa=2)


def _note(doc, text):
    _p(doc, f'（{text}）', size=9, color='666666', italic=True, sb=0, sa=2)


def _table(doc, headers, rows, col_widths=None,
           hdr_bg='1F3864', hdr_fg='FFFFFF', alt_bg='EEF3FB'):
    n_col = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_col)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 標題行
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, hdr_bg)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(str(h))
        _font(r, size=9.5, bold=True, color=hdr_fg)

    # 資料行
    for ri, row in enumerate(rows):
        bg = alt_bg if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            _set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val) if val is not None else '')
            _font(r, size=9.5)

    # 欄寬
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def _wallet_summary_table(doc, wallet_code: str, wallet_addr: str,
                           token_rows: list[list], extra_note=''):
    """生成錢包摘要表（幣種標題行+餘額/交易/時間/金額行）"""
    if extra_note:
        _p(doc, extra_note, size=9.5, color='555555', sb=0, sa=2)

    for token_block in token_rows:
        _table(doc,
               headers=['幣種', '餘額', '交易總次數',
                        '首次交易時間(UTC+8)', '最近交易時間(UTC+8)',
                        '收入總額', '支出總額', '收入次數', '支出次數'],
               rows=[token_block],
               col_widths=[2.2, 2.4, 2, 3.2, 3.2, 3, 3, 1.8, 1.8],
               hdr_bg='2F5496')

    _note(doc, '資料來源：https://dashboard.misttrack.io/')


# ─────────────────────────────────────────────────────────────────────────────
# 主函式
# ─────────────────────────────────────────────────────────────────────────────

def build_flow_report(data: dict, output_path: str) -> str:
    """
    依 data dict 生成幣流分析報告 Word 文件。
    data 空 dict → 輸出有完整提示文字的空白範本。
    """
    doc = Document()

    # 版面
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.0)

    doc.styles['Normal'].font.name = 'Microsoft JhengHei'
    doc.styles['Normal'].font.size = Pt(10.5)

    V = data

    # ═══════════════════════════════════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════════════════════════════════
    _p(doc, '幣流分析報告', bold=True, size=20, color='1F3864',
       align=WD_ALIGN_PARAGRAPH.CENTER, sb=16, sa=6)
    _p(doc, 'Cryptocurrency Flow Analysis Report',
       size=12, color='4472C4', align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)

    _table(doc,
           headers=['項目', '內容'],
           rows=[
               ['報告編號', V.get('report_number', '○○○字第○○○號')],
               ['案件名稱', V.get('case_name', '○○○詐欺案')],
               ['委託單位', V.get('client_unit', '臺北市政府警察局○○分局')],
               ['委託日期', V.get('commission_date', '')],
               ['報告日期', V.get('report_date', datetime.now().strftime('%Y年%m月%d日'))],
               ['鑑識分析人', V.get('analyst', '')],
               ['職稱/單位', V.get('analyst_title', '')],
               ['機密等級', V.get('classification', '限閱')],
           ],
           col_widths=[4, 13], hdr_bg='1F3864')

    doc.add_paragraph('─' * 68)

    # ═══════════════════════════════════════════════════════════════════════
    # 壹、案件背景與委託目的
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '壹、案件背景與委託目的')

    _h3(doc, '一、發生時間：')
    _p(doc, V.get('incident_time', '　○○年○○月間起至○○年○○月○○日'))

    _h3(doc, '二、發生地點：')
    _p(doc, V.get('incident_location', '　○○市○○區○○路○段○○號○樓之○'))

    _h3(doc, '三、發生經過：')
    _p(doc, V.get('incident_description',
        '　被害人○○○於○○年○○月間收到○○平台（帳號「○○○」暱稱「○○○」）交友私訊，'
        '被害人加入好友後，對方改提供LINE ID與被害人加好友（暱稱「○○○」），'
        '遭對方以交友及【○○話術】詐騙投資虛擬貨幣。'))

    # ═══════════════════════════════════════════════════════════════════════
    # 貳、資料來源與分析工具
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '貳、資料來源與分析工具')

    # 一、假投資平台
    _h2(doc, '一、依詐騙團提供之平臺網站：')
    fraud_sites = V.get('fraud_sites', [
        ['網址', '狀態', '分析結論'],
        ['https://○○○.top/○○○', '502 Bad Gateway（已下線）', '非官方金融網站'],
        ['https://○○○.ink/○○○', '無可辨識內容', '非官方金融網站'],
        ['https://○○○.com/○○○', '無可辨識內容', '非官方金融網站'],
    ])
    if fraud_sites and isinstance(fraud_sites[0], list) and not isinstance(fraud_sites[0][0], list):
        if len(fraud_sites) > 1:
            _table(doc, fraud_sites[0], fraud_sites[1:],
                   col_widths=[6, 4, 7.5], hdr_bg='7B2D8B', hdr_fg='FFFFFF')
    _p(doc,
       '　（一）、域名結構不尋常：使用非常見二級域名（如 .top、.ink），與官方金融機構域名格式不符。\n'
       '　（二）、無法正常開啟或看不到網站內容：可能已下線或根本不存在官方頁面。\n'
       '　（三）、無官方來源與信任資訊：不屬任何公認金融機構、銀行官方品牌或已登記之金融服務。',
       size=10, sa=4)
    _note(doc, f'依據被害人{V.get("victim_name","○○○")}警詢筆錄分析。')

    # 二、被害人錢包
    _h2(doc, '二、被害人提供之個人錢包地址：')
    victim_wallets = V.get('victim_wallets', [
        {
            'code': 'A01',
            'address': '0x○○○…○○○',
            'chain': 'ETH',
            'source': 'OKX 交易所托管錢包',
            'note': '被害人透過OKX交易所APP申請，入金後直接轉入交易所水庫（為交易所內部交易）。',
            'token_rows': [
                ['USDC-ERC20', '0.0 USDC', '○', '○', '○', '○,○○○ USDC', '○,○○○ USDC', '○', '○'],
            ],
        },
    ])
    for w in victim_wallets:
        _h3(doc, f"（{_roman_to_bracket(w.get('order',''))}）、{w['code']}：{w['address']}")
        _p(doc, f"　{w.get('note','')}", size=10)
        if w.get('token_rows'):
            _wallet_summary_table(doc, w['code'], w['address'], w['token_rows'])

    # 三、幣商錢包
    _h2(doc, '三、幣商錢包地址剖析：')
    dealer_wallets = V.get('dealer_wallets', [
        {
            'code': 'B01',
            'address': 'T○○○…○○○',
            'name': '○○幣商',
            'note': '非托管錢包，有USDT-TRC20及TRX交易紀錄。',
            'token_rows': [
                ['USDT-TRC20', '0.0 USDT', '○', '○', '○', '○,○○○ USDT', '○,○○○ USDT', '○', '○'],
            ],
        },
    ])
    for w in dealer_wallets:
        _h3(doc, f"（{_roman_to_bracket(w.get('order',''))}）、{w['code']}：{w['address']}")
        _p(doc, f"　{w.get('note','')}", size=10)
        if w.get('token_rows'):
            _wallet_summary_table(doc, w['code'], w['address'], w['token_rows'])

    # 四、詐團錢包
    _h2(doc, '四、詐欺集團指定錢包地址分析：')
    fraud_wallets = V.get('fraud_wallets', [
        {
            'code': 'C01',
            'address': 'T○○○…○○○',
            'note': '詐欺集團最終收款錢包，接收多個幣商轉入之USDT。',
            'token_rows': [
                ['USDT-TRC20', '0.0 USDT', '○', '○', '○', '○,○○○ USDT', '○,○○○ USDT', '○', '○'],
            ],
        },
    ])
    for w in fraud_wallets:
        _h3(doc, f"、{w['code']}：{w['address']}")
        _p(doc, f"　{w.get('note','')}", size=10)
        if w.get('token_rows'):
            _wallet_summary_table(doc, w['code'], w['address'], w['token_rows'])

    # 錢包代碼對照表
    _h2(doc, '錢包代碼與錢包地址對照表：')
    wallet_map = V.get('wallet_map', [
        ['A01', '○○○（被害人）', '被害人（OKX）', '被害人透過OKX申請'],
        ['A02', '○○○',         '被害人（Bitget）', '初步研判被害人不具掌控權'],
        ['A03', '○○○（被害人）', '被害人（OKX TRX）', '被害人透過OKX申請'],
        ['B01', '○○○',         '○○幣商',     '幣商錢包'],
        ['B02', '○○○',         '○○幣商',     '幣商錢包'],
        ['C01', '○○○',         '詐欺集團',   '詐團收款錢包'],
    ])
    _table(doc,
           headers=['代號', '錢包地址', '錢包來源', '附註'],
           rows=wallet_map,
           col_widths=[1.8, 8.2, 4, 3.5])

    # 五、交易紀錄
    _h2(doc, '五、交易紀錄：')

    face_to_face = V.get('face_to_face_transactions', [
        ['第○次', '○年○月○日 ○時○分',
         '○○市○○區○○路○段○號',
         'NT$○○○,○○○', '○○幣商', '○,○○○ USDT',
         '○○.○○ NT/USDT', '○○.○○ NT/USDT',
         '○○○→A○', ''],
    ])
    _h3(doc, '（一）現金面交購幣紀錄：')
    _table(doc,
           headers=['次', '日期時間', '地點', '現金(NT$)', '幣商',
                    'USDT數量', '成交匯率', '當日均價', '資金流向', '備註'],
           rows=face_to_face,
           col_widths=[1.2, 3.5, 4, 2.5, 2.5, 2.5, 2.5, 2.5, 3, 2.5],
           hdr_bg='385723', hdr_fg='FFFFFF')
    _note(doc, '資料來源：https://www.oklink.com/zh-hant 及 https://coinmarketcap.com/；均價計算方式為當日幣最高與最低之均價。')

    online_txs = V.get('online_transactions', [
        ['○', '○○年○月○日', '○○交易所',
         '綁定○○銀行 ○○○帳號', 'NT$○○,○○○', '○,○○○ USDC/USDT', '○○→A○→A○'],
    ])
    _h3(doc, '（二）線上購幣紀錄（交易所）：')
    _table(doc,
           headers=['次', '日期', '交易所', '綁定銀行/帳號', '法幣金額(NT$)', '購得數量', '資金流向'],
           rows=online_txs,
           col_widths=[1.2, 2.8, 3, 4.5, 3, 3.5, 3.5],
           hdr_bg='385723', hdr_fg='FFFFFF')

    # 損失統計
    _h3(doc, '（三）損失統計：')
    loss_summary = V.get('loss_summary', [
        ['現金面交', '○次', 'NT$○,○○○,○○○', ''],
        ['線上購幣', '○次', 'NT$○○○,○○○', ''],
        ['合計損失', '', 'NT$○,○○○,○○○', ''],
    ])
    _table(doc,
           headers=['類型', '次數', '金額(NT$)', '備註'],
           rows=loss_summary,
           col_widths=[4, 2, 5, 6.5],
           hdr_bg='7B3F00', hdr_fg='FFFFFF')

    # 六、金流圖位置
    _h2(doc, '六、金流圖：')
    _p(doc, '　（一）、A○-1：【請附加金流圖圖片】', size=10)
    _p(doc, '　（二）、A○-2：【請附加金流圖圖片】', size=10)
    _note(doc, '金流圖可由 CryptoAnalyzer 系統之「幣流關聯圖」功能生成後截圖插入。')

    # ═══════════════════════════════════════════════════════════════════════
    # 參、幣流事實與錢包交易紀錄
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '參、幣流事實與錢包交易紀錄')

    _h2(doc, '一、被害人錢包轉入紀錄：')
    inflow_rows = V.get('victim_inflow', [
        ['○', '○年○月○日 ○:○○', '○○○(B○○)', 'A○○', '○,○○○ USDT', '○○○(交易Hash)'],
    ])
    _table(doc,
           headers=['筆', '時間(UTC+8)', '來源錢包', '目標錢包', '金額', '交易Hash'],
           rows=inflow_rows,
           col_widths=[1.2, 3.5, 4, 4, 2.5, 7.5],
           hdr_bg='1F4E79', hdr_fg='FFFFFF')

    _h2(doc, '二、被害人錢包轉出至詐欺集團錢包紀錄（L0）：')
    l0_rows = V.get('l0_transfers', [
        ['○', '○年○月○日 ○:○○', 'A○○', 'C○○（詐團）', '○,○○○ USDT', '○○○'],
    ])
    _table(doc,
           headers=['筆', '時間(UTC+8)', '來源錢包', '目標錢包', '金額', '交易Hash'],
           rows=l0_rows,
           col_widths=[1.2, 3.5, 4, 5, 2.5, 7],
           hdr_bg='843C0C', hdr_fg='FFFFFF')

    _h2(doc, '三、詐欺集團錢包轉出至第二層紀錄（L1）：')
    l1_rows = V.get('l1_transfers', [
        ['○', '○年○月○日 ○:○○', 'C○○', 'L1_○○○', '○,○○○ USDT', '○○○'],
    ])
    _table(doc,
           headers=['筆', '時間(UTC+8)', '來源錢包', 'L1 目標錢包', '金額', '交易Hash'],
           rows=l1_rows,
           col_widths=[1.2, 3.5, 4, 5, 2.5, 7],
           hdr_bg='833122', hdr_fg='FFFFFF')
    _p(doc, '　（後續 L2、L3 層依同格式延伸，直至斷點。）', size=9.5, color='555555')

    _h2(doc, '四、案包錢包地址 TRX 來源分析：')
    _p(doc,
       '　說明詐欺集團錢包之 TRX（手續費）來源，用以判斷錢包所屬組織關聯性。\n'
       '　【請填入：TRX 來源錢包地址、轉入時間、金額、關聯說明】',
       size=10, color='555555', italic=True)

    _h2(doc, '五、幣流圖：')
    _p(doc, '　【請附上幣流關係示意圖（可由 CryptoAnalyzer 幣流關聯圖功能截圖）】',
       size=10, color='555555', italic=True)

    # ═══════════════════════════════════════════════════════════════════════
    # 肆、供述比對與交叉驗證分析
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '肆、供述比對與交叉驗證分析')

    supply_rows = V.get('statement_verification', [
        ['被害人陳述', '鏈上紀錄', '比對結果', '說明'],
        ['○年○月○日面交 NT$○○萬',
         '○年○月○日 ○:○○ B○→A○ ○,○○○ USDT',
         '✅ 吻合', '時間差在合理範圍（○分鐘內）'],
        ['交付現金後購得 USDT',
         '鏈上收款時間與金額一致',
         '✅ 吻合', '匯率略高於當日均價，符合幣商慣例'],
        ['後轉出至對方錢包',
         'A○→C○ 鏈上紀錄',
         '✅ 吻合', '轉出時間與供述一致'],
    ])
    _table(doc,
           headers=['被害人供述', '鏈上紀錄', '比對結果', '說明'],
           rows=[r for r in supply_rows[1:]],
           col_widths=[5, 5, 2.5, 5],
           hdr_bg='1F4E79', hdr_fg='FFFFFF')

    # ═══════════════════════════════════════════════════════════════════════
    # 伍、斷點分析與停止依據說明
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '伍、斷點分析與停止依據說明')

    breakpoint_rows = V.get('breakpoints', [
        ['層級', '錢包地址', '停止原因', '最後交易時間', '說明'],
        ['L1', '○○○', '□ 進入交易所\n□ 分散至多錢包\n□ 無後續交易',
         '○年○月○日', '【填入停止追蹤原因與依據】'],
    ])
    _table(doc,
           headers=['層級', '錢包地址', '停止原因', '最後交易時間', '說明'],
           rows=[r for r in breakpoint_rows[1:]],
           col_widths=[2, 5.5, 5, 3.5, 6.5])
    _p(doc,
       '　停止依據：依「虛擬貨幣金融犯罪調查實務指引」，追蹤至以下情形之一得停止：\n'
       '　①資金進入有KYC的交易所；②資金分散至100個以上錢包；'
       '③無法確認下一層資金歸屬；④超出本案委託分析範圍。',
       size=10, sa=4)

    # ═══════════════════════════════════════════════════════════════════════
    # 陸、錢包剖繪與歸屬分析
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '陸、錢包剖繪與歸屬分析')

    profiling_rows = V.get('wallet_profiling', [
        ['代號', '錢包地址', '角色', '控制方', '依據', '幣種', '總流量(USDT)', '活躍期間'],
        ['A01', '○○○', '被害人', '被害人',   'OKX KYC確認', 'USDC/ETH', '○○,○○○', '○～○'],
        ['A02', '○○○', '被害人？', '未確定', '非被害人申請，疑詐團控制', 'USDC', '○○,○○○', '○～○'],
        ['A03', '○○○', '被害人', '被害人',   'OKX KYC確認', 'USDT', '○○,○○○', '○～○'],
        ['B01', '○○○', '幣商',   '幣瑞',    '現金交付比對', 'USDT', '○○,○○○', '○～○'],
        ['B02', '○○○', '幣商',   'Honor Coin', '現金交付比對', 'USDT', '○○,○○○', '○～○'],
        ['C01', '○○○', '詐欺集團', '不明',   '接收多幣商轉入', 'USDT', '○○○,○○○', '○～○'],
    ])
    _table(doc,
           headers=profiling_rows[0],
           rows=profiling_rows[1:],
           col_widths=[1.8, 5, 2.5, 3, 4, 2, 3, 2.5])

    # ═══════════════════════════════════════════════════════════════════════
    # 柒（保留空白）
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '柒、其他補充事項（視需要填入）')
    _p(doc, '　【如有其他補充分析事項，請於此填寫。原稿無此章節，可刪除。】',
       size=10, color='999999', italic=True)

    # ═══════════════════════════════════════════════════════════════════════
    # 捌、分析結論與司法建議事項
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '捌、分析結論與司法建議事項')

    conclusions = V.get('conclusions', [
        '一、本案被害人○○○共遭詐騙新臺幣○,○○○,○○○元，資金流向確認如下：',
        '　（一）現金○次面交→幣商→詐欺集團錢包（C01），共○,○○○ USDT。',
        '　（二）線上購幣→被害人錢包（A01/A03）→詐欺集團錢包（C01）。',
        '二、詐欺集團錢包（C01）收款後，資金於○○年○○月間轉出，目前追蹤至【填入L1狀態】。',
        '三、幣商○○（B01）、○○（B02-B04）經核與面交時間及金額比對吻合，可確認其為本案資金中間人。',
    ])
    for c in conclusions:
        _p(doc, c, size=10.5)

    _h2(doc, '司法建議：')
    recommendations = V.get('recommendations', [
        '□ 一、凍結詐欺集團錢包（C01）及後續流向錢包，向各交易所申請帳戶KYC資料。',
        '□ 二、調取幣商【B01幣瑞、B02-B04 Honor Coin】相關帳冊及客戶資料，釐清幣商是否知情。',
        '□ 三、比對C01錢包與其他已知詐欺集團錢包，研判是否為同一組織運作。',
        '□ 四、向詐欺假投資平台之域名商申請域名持有人資料（透過司法協查）。',
        '□ 五、申請OKX及Bitget交易所提供A01、A02錢包KYC及交易紀錄。',
    ])
    for r in recommendations:
        _p(doc, r, size=10.5)

    # ═══════════════════════════════════════════════════════════════════════
    # 玖、附錄與附表（含鑑定人背景）
    # ═══════════════════════════════════════════════════════════════════════
    _h1(doc, '玖、附錄與附表（含鑑定人背景介紹）')

    _h2(doc, '一、鑑識分析人員背景：')
    analyst = V.get('analyst_background', {})
    _table(doc,
           headers=['項目', '內容'],
           rows=[
               ['姓名', analyst.get('name', '')],
               ['現職', analyst.get('current_position', '臺北市政府警察局○○分局偵查隊小隊長')],
               ['經歷', analyst.get('experience',
                   '刑事偵查工作○○年、曾獲國家警光獎詐欺類、全國績優刑事人員')],
               ['專業訓練', analyst.get('training',
                   'JM 金流分析實務工作坊（思偉達創新科技有限公司）\n'
                   '虛擬貨幣基礎原理及實務結業（睿科金融科技有限公司）')],
               ['分析工具', analyst.get('tools',
                   'MistTrack（https://dashboard.misttrack.io/）\n'
                   'OKLink（https://www.oklink.com/zh-hant）\n'
                   'CoinMarketCap（https://coinmarketcap.com/）\n'
                   'CryptoAnalyzer（自行開發鑑識系統）')],
           ],
           col_widths=[4, 13.5])

    _h2(doc, '二、附件清單：')
    attachments = V.get('attachments', [
        ['附件一', '被害人警詢筆錄', ''],
        ['附件二', 'OKX 錢包地址確認函（A01/A03）', ''],
        ['附件三', 'MistTrack 各錢包完整交易紀錄截圖', ''],
        ['附件四', '現金面交地點監視器截圖', ''],
        ['附件五', '假投資網站截圖', ''],
        ['附件六', '幣流關聯圖（完整版）', ''],
        ['附件七', '鑑識人員相關訓練結業證書', ''],
    ])
    _table(doc,
           headers=['編號', '附件名稱', '取得狀態'],
           rows=attachments,
           col_widths=[3, 10, 4.5])

    # 頁尾
    doc.add_paragraph('─' * 68)
    _p(doc,
       f"本報告由 CryptoAnalyzer 幣流分析系統輔助生成　{datetime.now().strftime('%Y-%m-%d %H:%M')}",
       size=9, color='999999', align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, '本報告含個人資料及偵查機密，請依法保管，未經授權不得複製或外洩。',
       size=9, color='CC0000', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    return output_path


def _roman_to_bracket(s):
    """回傳中文括號序號，或原字串"""
    m = {'1': '一', '2': '二', '3': '三', '4': '四', '5': '五',
         '6': '六', '7': '七', '8': '八', '9': '九'}
    return m.get(str(s), str(s))


# ─────────────────────────────────────────────────────────────────────────────
# 黃瑋琳案完整資料（直接從文件提取）
# ─────────────────────────────────────────────────────────────────────────────

HUANG_CASE_DATA = {
    'report_number':   '松山分局○○字第○○○號',
    'case_name':       '黃瑋琳 IG 虛擬貨幣投資詐欺案',
    'client_unit':     '臺北市政府警察局松山分局',
    'commission_date': '114年12月25日',
    'report_date':     '115年○○月○○日',
    'analyst':         '吳家宵',
    'analyst_title':   '臺北市政府警察局松山分局偵查隊小隊長',
    'classification':  '限閱',

    # 壹
    'incident_time':        '114年9月間起至114年12月25日',
    'incident_location':    '臺北市松山區八德路2段348號5樓之2（被害人住所）',
    'victim_name':          '黃瑋琳',
    'incident_description': (
        '　被害人黃瑋琳於114年9月間收到IG帳號「kiraa.1u4」暱稱「Fen_1031」交友私訊，'
        '被害人加入好友後，對方改提供LINE ID（ID不詳）與被害人加好友，好友暱稱為「陳裕峰」，'
        '之後被害人自行將「陳裕峰」之暱稱改為「AB型天蠍座♏」，'
        '遭對方以交友及共組家庭之話術詐騙投資虛擬貨幣。'
    ),

    # 貳：假投資平台
    'fraud_sites': [
        ['假投資網站 URL', '狀態', '分析結論'],
        ['https://tornbanka.top/1d55eb28',   '502 Bad Gateway（已下線）', '非官方金融網站，已無法存取'],
        ['https://torninbank.ink/bb05f247', '無可辨識內容',              '非官方金融網站，無標準銀行頁面'],
        ['https://imtornbank.com/f1a538b3', '無可辨識內容',              '非官方金融網站，無標準銀行頁面'],
    ],

    # 被害人錢包
    'victim_wallets': [
        {
            'order': '1', 'code': 'A01',
            'address': '0x3F3b240D6cBbD69DF3126772b7Fb592ec6a55f6a',
            'chain': 'ETH',
            'note': '該錢包為OKX交易所托管錢包，有USDC-ERC20及ETH交易紀錄。'
                    '經OKX公司查覆，確為被害人黃瑋琳所申請使用（詳附件二），'
                    '入金後均直接轉入交易所水庫錢包（交易所內部交易）。',
            'token_rows': [
                ['USDC-ERC20', '0.0 USDC', '26', '2025/11/3 12:51', '2025/12/11 17:26',
                 '16,307.6294 USDC', '16,307.6294 USDC', '13', '13'],
                ['ETH', '0.0006 ETH', '2', '2025/11/3 12:51', '2025/11/5 14:04',
                 '0.0008 ETH', '0.0 ETH', '2', '0'],
            ],
        },
        {
            'order': '2', 'code': 'A02',
            'address': '0x978309e037dc0a94733232A0CD254aB0A9c00Feb',
            'chain': 'ETH',
            'note': '該錢包地址為Bitget wallet APP所產生，有USDC-ERC20及ETH交易紀錄。'
                    'ETH第一次交易紀錄在1970年（Unix Epoch起始值），USDC第一次交易在2025年9月25日，'
                    '此二交易與被害人均無關聯，可確認該錢包應非被害人所申請使用，'
                    '被害人初步研判不具掌控權限。',
            'token_rows': [
                ['USDC-ERC20', '0.0124 USDC', '34', '2025/9/25 21:27', '2025/12/11 20:41',
                 '58,728.8124 USDC', '58,728.8 USDC', '28', '6'],
                ['ETH', '0.0 ETH', '0', '1970/1/1 08:00', '1970/1/1 08:00',
                 '0.0 ETH', '0.0 ETH', '0', '0'],
            ],
        },
        {
            'order': '3', 'code': 'A03',
            'address': 'TY4sP2pVz4oCwiDkGmX5ubreeQB1mGpd4B',
            'chain': 'TRX',
            'note': '該錢包為OKX交易所托管錢包，有USDT-TRC20交易紀錄。'
                    '經OKX公司查覆，確為被害人黃瑋琳所申請使用（詳附件二），'
                    '入金後均直接轉入交易所水庫（交易所內部交易）。',
            'token_rows': [
                ['USDT-TRC20', '0.0 USDT', '9', '2025/11/5 21:25', '2025/11/10 19:32',
                 '39,519.0 USDT', '39,519.0 USDT', '6', '3'],
            ],
        },
    ],

    # 幣商錢包
    'dealer_wallets': [
        {
            'order': '1', 'code': 'B01', 'name': '幣瑞',
            'address': 'TSG6UX57c2c9kr7StuKMW4eRftbD1fgAJn',
            'note': '非托管錢包，有USDT-TRC20及TRX交易紀錄。（幣商：幣瑞）',
            'token_rows': [
                ['USDT-TRC20', '0.0 USDT', '513', '2025/10/30 20:04', '2025/11/29 00:11',
                 '4,796,202.9073 USDT', '4,796,202.9073 USDT', '216', '297'],
                ['TRX', '1.9092 TRX', '2', '2025/10/30 23:17', '2025/11/5 10:14',
                 '60.9647 TRX', '0.0 TRX', '2', '0'],
            ],
        },
        {
            'order': '2', 'code': 'B02', 'name': 'Honor Coin',
            'address': 'TP2d171RRNmg5GTZq24UzjmNCPUoC7DnwD',
            'note': '非托管錢包，有USDT-TRC20及TRX交易紀錄。（幣商：Honor Coin）',
            'token_rows': [
                ['USDT-TRC20', '30,563.6628 USDT', '1,524', '2025/9/12 22:56', '2026/1/3 11:10',
                 '9,079,859.9968 USDT', '9,049,296.334 USDT', '425', '1,099'],
                ['TRX', '203.054 TRX', '18', '2025/9/12 22:57', '2025/12/15 10:25',
                 '646.002 TRX', '121.872 TRX', '5', '13'],
            ],
        },
        {
            'order': '3', 'code': 'B03', 'name': 'Honor Coin',
            'address': 'TVe7uMmtta7rhJUyrjV9aEhbjNWzwgLRz8',
            'note': '非托管錢包，有USDT-TRC20及TRX交易紀錄。（幣商：Honor Coin）',
            'token_rows': [
                ['USDT-TRC20', '0.0 USDT', '609', '2025/9/15 12:09', '2025/11/25 21:55',
                 '3,790,494.221 USDT', '3,790,494.221 USDT', '292', '317'],
                ['TRX', '0.269 TRX', '28', '2025/9/14 22:46', '2025/11/25 23:29',
                 '120.0751 TRX', '88.3618 TRX', '20', '8'],
            ],
        },
        {
            'order': '4', 'code': 'B04', 'name': 'Honor Coin',
            'address': 'TPbBtJTYekVxqaBX5FrFzYMUgKeeiuUQbe',
            'note': '非托管錢包，有USDT-TRC20及TRX交易紀錄。（幣商：Honor Coin）',
            'token_rows': [
                ['USDT-TRC20', '0.0 USDT', '535', '2025/11/10 09:56', '2025/11/20 16:45',
                 '3,731,523.154 USDT', '3,731,523.154 USDT', '164', '371'],
                ['TRX', '0.0 TRX', '29', '2025/11/8 20:10', '2025/11/20 16:45',
                 '100.4385 TRX', '81.4991 TRX', '25', '4'],
            ],
        },
    ],

    # 詐團錢包
    'fraud_wallets': [
        {
            'code': 'C01',
            'address': 'THsAZgi1Hr9sqF9cKGTJgHUkwpo52qZvGz',
            'note': '詐欺集團指定最終收款錢包，接收多名幣商轉入之USDT，後續流向待追蹤。',
            'token_rows': [
                ['USDT-TRC20', '4,601.0 USDT', '262', '2025/10/11 17:42', '2025/11/29 13:50',
                 '1,645,437.6314 USDT', '1,640,836.6314 USDT', '99', '163'],
                ['TRX', '0.0 TRX', '42', '2025/11/4 20:27', '2025/11/28 18:09',
                 '324.982 TRX', '196.4417 TRX', '6', '36'],
            ],
        },
    ],

    # 錢包對照表
    'wallet_map': [
        ['A01', '0x3F3b240D6cBbD69DF3126772b7Fb592ec6a55f6a', '被害人（OKX ERC20）', '被害人透過OKX交易所APP申請'],
        ['A02', '0x978309e037dc0a94733232A0CD254aB0A9c00Feb', '被害人（Bitget）',     '初步研判被害人不具掌控權限'],
        ['A03', 'TY4sP2pVz4oCwiDkGmX5ubreeQB1mGpd4B',       '被害人（OKX TRC20）', '被害人透過OKX交易所APP申請'],
        ['B01', 'TSG6UX57c2c9kr7StuKMW4eRftbD1fgAJn',       '幣瑞（幣商）',        '幣商非托管錢包'],
        ['B02', 'TP2d171RRNmg5GTZq24UzjmNCPUoC7DnwD',       'Honor Coin（幣商）',  '幣商非托管錢包'],
        ['B03', 'TVe7uMmtta7rhJUyrjV9aEhbjNWzwgLRz8',       'Honor Coin（幣商）',  '幣商非托管錢包'],
        ['B04', 'TPbBtJTYekVxqaBX5FrFzYMUgKeeiuUQbe',       'Honor Coin（幣商）',  '幣商非托管錢包'],
        ['C01', 'THsAZgi1Hr9sqF9cKGTJgHUkwpo52qZvGz',       '詐欺集團',           '最終收款錢包'],
    ],

    # 現金面交
    'face_to_face_transactions': [
        ['第1次', '2025/11/5 21:28', '臺北市松山區八德路2段348號1樓大廳',
         '297,000', '幣瑞(B01)', '9,281 USDT', '32.00', '30.95', 'B01→A03', ''],
        ['第2次', '2025/11/7 12:47', '松山區八德路2段360號對面→長安東路2段259號（車內）',
         '630,000', 'Honor Coin(B02)', '19,385 USDT', '32.499', '30.985', 'B02→A03', ''],
        ['第3次', '2025/11/10 19:16', '臺安醫院前（八德路2段424號）車上',
         '350,000', 'Honor Coin(B03)', '10,853 USDT', '32.249', '30.985', 'B03→A03', ''],
        ['第4次', '2025/11/12 12:03', '思味咖啡店（八德路2段352號）',
         '770,000', 'Honor Coin(B03)', '23,692 USDT', '32.50', '31.045', 'B03→C01', '幣商直轉詐團錢包'],
        ['第5次', '2025/11/14 13:28', '台北市中正區羅斯福路二段102號旁巷子',
         '600,000', 'Honor Coin(B04)', '18,462 USDT', '32.499', '30.895', 'B04→C01', '幣商直轉詐團錢包'],
        ['第6次', '2025/11/17 21:05', '臺北市松山區八德路2段348號1樓',
         '240,000', 'Honor Coin(B04)', '7,407 USDT', '32.40', '31.265', 'B04→C01', '幣商直轉詐團錢包'],
        ['第7次', '2025/11/20 18:00', '臺北市松山區八德路二段348號外黑色自小客車上',
         '400,000', 'Honor Coin(B03)', '12,158 USDT', '32.90', '30.92', 'B03→C01', '幣商直轉詐團錢包'],
        ['第8次', '2025/11/25 10:57', '臺北市松山區八德路2段350號前黑色車輛內',
         '450,000', 'Honor Coin(B03)', '13,554 USDT', '33.20', '31.425', 'B03→C01', '幣商直轉詐團錢包'],
        ['小計', '', '', '3,737,000', '', '115,812 USDT', '', '', '', ''],
    ],

    # 線上購幣
    'online_transactions': [
        ['1', '114/11/3,4,11\n各 NT$30,000', 'Maicoin',
         '台新銀行 20111500016770', 'NT$90,000（3次）',
         '965.13+962.71+956.92=2,884.76 USDC', 'A01→A02'],
        ['2', '114/11/3~9\n多次', 'Bitopro',
         '第一銀行 16357026615', 'NT$340,000（7次）',
         '3,056.19+3,050.24+960.69+962.09+961.47+961.47+955.60=11,907.75 USDC', 'A01→A02'],
        ['小計', '', '', '', 'NT$430,000', '≈14,792 USDC', ''],
    ],

    # 損失統計
    'loss_summary': [
        ['現金面交（8次）', '8次', 'NT$3,737,000', '透過4個幣商錢包中轉'],
        ['線上購幣', '10次', 'NT$430,000', 'Maicoin+Bitopro，A01→A02'],
        ['合計損失', '', 'NT$4,167,000', '（約合115,812+14,792≈130,604 USDT）'],
    ],

    # 分析結論
    'conclusions': [
        '一、被害人黃瑋琳遭詐騙損失共計新臺幣4,167,000元，資金流向如下：',
        '　（一）現金面交8次，共NT$3,737,000，透過幣商（幣瑞/Honor Coin）購得共計115,812 USDT：',
        '　　　● 第1-3次：B01/B02/B03→被害人A03錢包（OKX托管），再由OKX內部轉至C01。',
        '　　　● 第4-8次：B03/B04直接轉至詐欺集團錢包C01，被害人A03未過帳。',
        '　（二）線上購幣10次，共NT$430,000（Maicoin+Bitopro），購得約14,792 USDC至A01，再轉A02。',
        '二、詐欺集團最終收款錢包C01（THsAZgi1Hr9sqF9cKGTJgHUkwpo52qZvGz）總收款逾164萬USDT，顯示本案為大規模詐欺集團。',
        '三、幣商B02（Honor Coin：TP2d...）單錢包累積流量逾900萬USDT，活躍期間2025/9-2026/1，高度懷疑為職業非法幣商。',
    ],

    'recommendations': [
        '□ 一、向TronScan/OKLink申請C01錢包後續流向完整紀錄，追蹤L1/L2層。',
        '□ 二、凍結及扣押B01-B04幣商非托管錢包（TRX鏈），偵查幣商身份。',
        '□ 三、調取面交地點（八德路2段348/350/360號等地）監視器影像，比對幣商外貌。',
        '□ 四、申請OKX提供A01/A03錢包完整KYC與交易紀錄（A01被害人確認所有；A03被害人所有）。',
        '□ 五、申請Bitget提供A02錢包KYC，確認實際掌控人（疑非被害人）。',
        '□ 六、比對C01與已知詐欺集團錢包資料庫，研判是否跨案關聯。',
        '□ 七、就假投資平台域名（tornbanka.top等）向TWNIC或國際域名商申請申請人資料。',
    ],

    'analyst_background': {
        'name':             '吳家宵',
        'current_position': '臺北市政府警察局松山分局偵查隊小隊長',
        'experience':       '刑事偵查工作11年、曾獲國家警光獎詐欺類、全國績優刑事人員',
        'training': (
            'JM 金流分析實務工作坊（思偉達創新科技有限公司）\n'
            '虛擬貨幣基礎原理及實務結業（睿科金融科技有限公司）'
        ),
        'tools': (
            'MistTrack（https://dashboard.misttrack.io/）\n'
            'OKLink（https://www.oklink.com/zh-hant）\n'
            'CoinMarketCap（https://coinmarketcap.com/）'
        ),
    },
}


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import sys
    mode = sys.argv[1] if len(sys.argv) > 1 else 'case'
    base = r'D:\claude\crypto_analyzer'

    if mode == 'template':
        path = os.path.join(base, '幣流分析報告_通用範本_空白.docx')
        build_flow_report({}, path)
        print(f'空白範本：{path}')
    else:
        path = os.path.join(base, '幣流分析報告_黃瑋琳案.docx')
        build_flow_report(HUANG_CASE_DATA, path)
        print(f'本案報告：{path}')
