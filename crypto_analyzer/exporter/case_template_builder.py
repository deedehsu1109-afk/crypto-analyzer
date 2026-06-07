"""
case_template_builder.py
生成「虛擬貨幣詐欺案件分析範本」Word 文件。
支援：空白範本 / 填入案件資料後匯出。
"""
from __future__ import annotations
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# 樣式輔助
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """設定儲存格背景色，hex_color 例如 '1F3864'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell):
    """設定細框線"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), '4')
        bd.set(qn('w:color'), '999999')
        tcBorders.append(bd)
    tcPr.append(tcBorders)


def _para(doc, text='', bold=False, size=11, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Microsoft JhengHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return p


def _heading(doc, text: str, level=1):
    colors = {1: '1F3864', 2: '2E75B6', 3: '404040'}
    sizes  = {1: 14, 2: 12, 3: 11}
    _para(doc, text, bold=True, size=sizes.get(level, 11),
          color=colors.get(level, '000000'),
          space_before=8, space_after=4)


def _add_table(doc, headers: list[str], rows: list[list],
               col_widths: list[float] | None = None,
               header_bg: str = '1F3864',
               header_fg: str = 'FFFFFF') -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 標題行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, header_bg)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Microsoft JhengHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
        run.font.color.rgb = RGBColor.from_string(header_fg)

    # 資料行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'F5F8FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val) if val is not None else '')
            run.font.size = Pt(10)
            run.font.name = 'Microsoft JhengHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # 欄寬
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# 主建構函式
# ─────────────────────────────────────────────────────────────────────────────

def build_case_doc(data: dict, output_path: str) -> str:
    """
    依 data dict 生成案件分析 Word 文件並存檔。
    data 可為空 dict（生成空白範本）或填入實際案件資料。
    """
    doc = Document()

    # ── 版面設定 ──
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    # ── 預設字型 ──
    doc.styles['Normal'].font.name = 'Microsoft JhengHei'
    doc.styles['Normal'].font.size = Pt(10.5)

    V = data  # shorthand

    # ═══════════════════════════════════════════════════════════════════════
    # 封面標題
    # ═══════════════════════════════════════════════════════════════════════
    _para(doc, '虛擬貨幣詐欺案件分析報告', bold=True, size=18,
          color='1F3864', align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=12, space_after=4)
    _para(doc, 'Cryptocurrency Fraud Case Analysis Report',
          bold=False, size=11, color='4472C4',
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, f"製作日期：{V.get('report_date', datetime.now().strftime('%Y-%m-%d'))}　"
               f"製作人：{V.get('analyst', '__________')}　"
               f"機密等級：{V.get('classification', '限閱')}",
          size=10, color='666666', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    doc.add_paragraph('─' * 60)

    # ═══════════════════════════════════════════════════════════════════════
    # 一、案件基本資料
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '一、案件基本資料', level=1)
    _add_table(doc,
        headers=['項目', '內容'],
        rows=[
            ['案件編號',    V.get('case_number', '')],
            ['案件名稱',    V.get('case_name', '')],
            ['案件類型',    V.get('case_type', '虛擬貨幣詐欺（豬仔盤）')],
            ['承辦單位',    V.get('unit', '')],
            ['承辦人員',    V.get('investigator', '')],
            ['報案日期',    V.get('report_date_incident', '')],
            ['製作筆錄日期', V.get('record_date', '')],
            ['筆錄地點',    V.get('record_location', '')],
            ['詢問時間',    V.get('inquiry_time', '')],
            ['詢問人員',    V.get('inquiry_officer', '')],
            ['案件狀態',    V.get('status', '偵查中')],
        ],
        col_widths=[5, 12])

    # ═══════════════════════════════════════════════════════════════════════
    # 二、被害人基本資料
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '二、被害人基本資料', level=1)
    _add_table(doc,
        headers=['項目', '內容', '備註'],
        rows=[
            ['姓名',       V.get('victim_name', ''),          ''],
            ['性別',       V.get('victim_gender', ''),        ''],
            ['出生年月日', V.get('victim_dob', ''),           '民國年或西元年'],
            ['身分證字號', V.get('victim_id', ''),            '司法保密'],
            ['職業',       V.get('victim_occupation', ''),    ''],
            ['教育程度',   V.get('victim_education', ''),     ''],
            ['戶籍地址',   V.get('victim_address_reg', ''),   ''],
            ['現住地址',   V.get('victim_address_cur', ''),   ''],
            ['聯絡電話',   V.get('victim_phone', ''),         ''],
            ['電子郵件',   V.get('victim_email', ''),         ''],
            ['家庭經濟',   V.get('victim_economy', ''),       ''],
            ['資金來源',   V.get('victim_fund_source', ''),   '自有存款/貸款/借款'],
        ],
        col_widths=[4, 8, 5])

    # ═══════════════════════════════════════════════════════════════════════
    # 三、詐騙手法與接觸過程
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '三、詐騙手法與接觸過程', level=1)
    _add_table(doc,
        headers=['項目', '內容'],
        rows=[
            ['初始接觸平台',  V.get('contact_platform', '')],
            ['初始接觸日期',  V.get('contact_date', '')],
            ['詐騙手法分類',  V.get('fraud_method', '')],
            ['假投資網站 URL', V.get('fraud_website', '')],
            ['假APP下載連結', V.get('fraud_app', '無')],
            ['假冒機構名稱',  V.get('fake_org', '無')],
            ['通訊方式',      V.get('communication', '')],
            ['面交地點',      V.get('meeting_location', '')],
            ['面交次數',      V.get('meeting_count', '')],
            ['面交人員描述',  V.get('meeting_person_desc', '')],
        ],
        col_widths=[5, 12])

    # ═══════════════════════════════════════════════════════════════════════
    # 四、嫌疑人帳號資訊
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '四、嫌疑人帳號資訊', level=1)
    suspect_rows = V.get('suspect_accounts', [
        ['LINE', '', ''],
        ['Instagram', '', ''],
        ['Telegram', '', ''],
        ['Facebook', '', ''],
        ['其他平台', '', ''],
    ])
    _add_table(doc,
        headers=['平台', '帳號/ID', '備註'],
        rows=suspect_rows,
        col_widths=[4, 9, 4])

    # ═══════════════════════════════════════════════════════════════════════
    # 五、資金損失明細
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '五、資金損失明細', level=1)

    _heading(doc, '5-1 銀行轉帳紀錄', level=2)
    bank_rows = V.get('bank_transfers', [
        ['', '', '', '', '', ''],
    ])
    _add_table(doc,
        headers=['筆次', '日期時間', '匯款帳戶', '收款帳號（銀行）', '金額(NT$)', '備註'],
        rows=bank_rows,
        col_widths=[1.5, 3.5, 4.5, 4.5, 2.5, 2])

    _heading(doc, '5-2 現金交付紀錄', level=2)
    cash_rows = V.get('cash_payments', [
        ['', '', '', '', ''],
    ])
    _add_table(doc,
        headers=['次', '日期時間', '地點', '金額(NT$)', '收款描述/備註'],
        rows=cash_rows,
        col_widths=[1.5, 3.5, 4, 2.5, 6.5])

    _heading(doc, '5-3 虛擬貨幣轉帳紀錄', level=2)
    crypto_rows = V.get('crypto_transfers', [
        ['', '', '', '', '', ''],
    ])
    _add_table(doc,
        headers=['筆次', '日期時間', '幣種', '數量', '被害人錢包（FROM）', '嫌疑人錢包（TO）'],
        rows=crypto_rows,
        col_widths=[1.5, 3.5, 2, 2.5, 5.5, 5.5])

    _heading(doc, '5-4 損失總計', level=2)
    total_rows = V.get('total_loss', [
        ['銀行轉帳', '', ''],
        ['現金交付', '', ''],
        ['虛擬貨幣（換算NT$）', '', ''],
        ['其他', '', ''],
        ['合計', '', ''],
    ])
    _add_table(doc,
        headers=['類別', '金額(NT$)', '備註'],
        rows=total_rows,
        col_widths=[6, 4, 8])

    # ═══════════════════════════════════════════════════════════════════════
    # 六、區塊鏈錢包分析
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '六、區塊鏈錢包分析', level=1)
    wallet_rows = V.get('wallets', [
        ['被害人', 'TRX', '', ''],
        ['嫌疑人', 'TRX', '', ''],
        ['中間人/幣商', 'TRX', '', ''],
    ])
    _add_table(doc,
        headers=['角色', '鏈', '錢包地址', '說明'],
        rows=wallet_rows,
        col_widths=[3, 2, 9, 4])

    _heading(doc, '分析建議', level=2)
    suggestions = V.get('wallet_analysis_notes', [
        '□ 嫌疑人錢包上鏈查詢（TRX 鏈）',
        '□ 追蹤嫌疑人錢包後續資金流向',
        '□ 確認幣商中介錢包地址',
        '□ 查詢嫌疑人收款帳號銀行開戶資料',
        '□ 申請交易所（BingX）帳號KYC資料',
    ])
    for s in suggestions:
        _para(doc, s, size=10.5)

    # ═══════════════════════════════════════════════════════════════════════
    # 七、事件時序
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '七、事件時序（Timeline）', level=1)
    timeline_rows = V.get('timeline', [
        ['', '', '', ''],
    ])
    _add_table(doc,
        headers=['日期', '時間', '事件', '說明/來源'],
        rows=timeline_rows,
        col_widths=[3.5, 2.5, 6, 6])

    # ═══════════════════════════════════════════════════════════════════════
    # 八、證據清單
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '八、證據清單', level=1)
    evidence_rows = V.get('evidence', [
        ['E-001', '調查筆錄', '紙本', '', '○'],
        ['E-002', '轉帳紀錄（ATM/網銀截圖）', '電子', '', '□'],
        ['E-003', '假投資網站截圖', '電子', '', '□'],
        ['E-004', 'LINE/IG 對話截圖', '電子', '', '□'],
        ['E-005', '現金交付收據/幣商收據', '紙本', '', '□'],
        ['E-006', '虛擬貨幣轉帳紀錄（區塊鏈）', '電子', '', '□'],
        ['E-007', '銀行帳號開戶資料（公文調取）', '紙本', '', '□'],
        ['E-008', 'BingX帳號KYC資料（公文調取）', '電子', '', '□'],
    ])
    _add_table(doc,
        headers=['編號', '證據名稱', '類型', '存放位置', '已取得'],
        rows=evidence_rows,
        col_widths=[2, 7, 2.5, 4.5, 2])

    # ═══════════════════════════════════════════════════════════════════════
    # 九、分析摘要與偵查建議
    # ═══════════════════════════════════════════════════════════════════════
    _heading(doc, '九、分析摘要與偵查建議', level=1)
    _heading(doc, '手法特徵', level=2)
    features = V.get('fraud_features', [
        '• 「殺豬盤」手法：建立信任關係後誘導投資假平台',
        '• 透過合法交友App（BUMBLE）接觸被害人',
        '• 使用幣商OTC現金交易，規避銀行KYC',
        '• 以虛擬貨幣（USDT/TRX鏈）作為資金轉移媒介',
        '• 假投資平台可正常顯示「獲利」，誘使持續匯款',
        '• 最終以「提款需繳稅/保證金」手法拒絕出金',
    ])
    for f in features:
        _para(doc, f, size=10.5)

    _heading(doc, '偵查建議', level=2)
    recommendations = V.get('recommendations', [
        '1. 凍結對方收款帳號（004-031001010808），向銀行申請帳戶資料',
        '2. 向 TRonScan 查詢嫌疑人錢包 TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL 後續流向',
        '3. 調取幣商營業地（八德路三段200號）監視器影像',
        '4. 協請 BingX 交易所提供 KYC 帳號資料（司法協查）',
        '5. 調查假投資網站 www.manareszke.com 主機 IP 及域名登記資料',
        '6. 比對其他案件中相同錢包地址或帳號，研判是否為集團犯罪',
    ])
    for r in recommendations:
        _para(doc, r, size=10.5)

    # ═══════════════════════════════════════════════════════════════════════
    # 頁尾
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_paragraph('─' * 60)
    _para(doc, f"本文件由 CryptoAnalyzer 系統自動生成　生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
          size=9, color='999999', align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, '本文件含個人資料及偵查資訊，請依法保管，未經授權不得外洩。',
          size=9, color='CC0000', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    # ── 存檔 ──
    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# 本案實際資料（1-3.pdf）
# ─────────────────────────────────────────────────────────────────────────────

CASE_1_3_DATA = {
    # 報告基本
    'report_date':       datetime.now().strftime('%Y-%m-%d'),
    'analyst':           '吳家宵',
    'classification':    '限閱',

    # 案件
    'case_number':       'CASE-20250111-001',
    'case_name':         '陳姿穎 BUMBLE 虛擬貨幣詐欺案',
    'case_type':         '虛擬貨幣詐欺（殺豬盤）',
    'unit':              '臺北市政府警察局松山分局',
    'investigator':      '吳家宵',
    'report_date_incident': '115年01月11日',
    'record_date':       '115年01月11日',
    'record_location':   '臺北市松山區南京東路四段12號',
    'inquiry_time':      '08:08 ～ 09:43',
    'inquiry_officer':   '警員吳家宵',
    'status':            '偵查中',

    # 被害人
    'victim_name':           '陳姿穎',
    'victim_gender':         '女',
    'victim_dob':            '民國088年09月04日（1999-09-04）',
    'victim_id':             'A230146012',
    'victim_occupation':     '醫療護理藥劑專業人員',
    'victim_education':      '大學畢業',
    'victim_address_reg':    '桃園市桃園區大連四街18號8樓',
    'victim_address_cur':    '臺北市松山區光復南路6巷47號4樓',
    'victim_phone':          '0912585555',
    'victim_email':          '',
    'victim_economy':        '小康',
    'victim_fund_source':    '自有存款＋國泰世華信用貸款 NT$250,000',

    # 詐騙
    'contact_platform':  'BUMBLE（交友App）',
    'contact_date':      '114年12月29日（2025-12-29）',
    'fraud_method':      '殺豬盤 / 假投資虛擬貨幣',
    'fraud_website':     'www.manareszke.com',
    'fraud_app':         '無',
    'fake_org':          '無',
    'communication':     'LINE（無電話）',
    'meeting_location':  '臺北市松山區八德路三段200號',
    'meeting_count':     '2次',
    'meeting_person_desc': (
        '第1次：微胖短髮背後背包中年女性；'
        '第2次：咖啡色長髮馬尾中年女性（均無車牌）'
    ),

    # 嫌疑人帳號
    'suspect_accounts': [
        ['LINE',      'hao._07',   '交友網站認識後轉至LINE聯繫'],
        ['Instagram', 'gary11.__', ''],
        ['其他',      '',          ''],
    ],

    # 銀行轉帳
    'bank_transfers': [
        ['第1筆', '115-01-01 21:04', '013-69951386361', '004-031001010808', '3,000', ''],
        ['第2筆', '115-01-01 21:57', '824-111080083398', '004-031001010808', '50,000', ''],
        ['小計',  '',                '',               '',                  '53,000', ''],
    ],

    # 現金
    'cash_payments': [
        ['第1次', '115-01-07 17時許', '台北市松山區八德路三段200號', '300,000', '幣商收現金後直轉USDT至對方錢包'],
        ['第2次', '115-01-09 17時許', '台北市松山區八德路三段200號', '450,000', '幣商轉USDT進被害人BingX帳戶'],
        ['小計',  '',                  '',                           '750,000', ''],
    ],

    # 虛擬貨幣轉帳
    'crypto_transfers': [
        ['第1筆', '115-01-01 23:12', 'USDT(TRX)', '3,078.50',
         'TPpSCUoTFosfujkDGjdrSEXspTpTQfTQvy',
         'TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL'],
        ['第2筆', '115-01-02 18:45', 'USDT(TRX)', '3,080.50',
         'TPpSCUoTFosfujkDGjdrSEXspTpTQfTQvy',
         'TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL'],
        ['第3筆', '115-01-02',       'USDT(TRX)', '3,082.50',
         'TPpSCUoTFosfujkDGjdrSEXspTpTQfTQvy',
         'TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL'],
        ['第4筆', '115-01-02',       'USDT(TRX)', '4,624.19',
         'TPpSCUoTFosfujkDGjdrSEXspTpTQfTQvy',
         'TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL'],
        ['第5筆', '115-01-07',       'USDT(TRX)', '557.00',
         'TPpSCUoTFosfujkDGjdrSEXspTpTQfTQvy',
         'TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL'],
        ['第6筆', '115-01-09',       'USDT(TRX)', '12,610.00',
         'TPpSCUoTFosfujkDGjdrSEXspTpTQfTQvy',
         'TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL'],
        ['合計',  '',                '',          '27,032.69 USDT', '', ''],
    ],

    # 損失總計
    'total_loss': [
        ['銀行轉帳',        '53,000',    '2筆，匯至004-031001010808'],
        ['現金交付',        '750,000',   '2次面交，八德路三段200號幣商'],
        ['其他（差額）',    '418,159',   '文件記載總額-已知明細，待查'],
        ['文件記載損失合計', '1,221,159', '調查筆錄第2頁記載'],
        ['※ 注意',         '',          '使用者摘要稱NT$1,121,159，與文件差NT$100,000'],
    ],

    # 錢包
    'wallets': [
        ['被害人', 'TRX(USDT)', 'TPpSCUoTFosfujkDGjdrSEXspTpTQfTQvy', '陳姿穎持有'],
        ['嫌疑人', 'TRX(USDT)', 'TTktL5ebm63sU1BNY5wGGvg8mpDf2tsGL', '詐騙收款'],
        ['交易所', 'TRX',       'BingX 帳戶（KYC待查）',               '被害人開立，幣商轉入後轉出'],
    ],

    # 時序
    'timeline': [
        ['114-12-29', '',        '初始接觸', 'BUMBLE 認識網友（LINE:hao._07）'],
        ['115-01-01', '21:04',   '銀行轉帳', 'NT$3,000 → 0808帳號'],
        ['115-01-01', '21:57',   '銀行轉帳', 'NT$50,000 → 0808帳號'],
        ['115-01-01', '23:12',   'USDT轉帳', '3,078.5 USDT → 嫌疑人錢包'],
        ['115-01-02', '18:45',   'USDT轉帳', '3,080.5 USDT → 嫌疑人錢包'],
        ['115-01-02', '',        'USDT轉帳', '3,082.5 + 4,624.19 USDT → 嫌疑人錢包'],
        ['115-01-07', '17時許', '現金面交', 'NT$300,000，八德路三段200號'],
        ['115-01-07', '',        'USDT轉帳', '557 USDT → 嫌疑人錢包'],
        ['115-01-09', '17時許', '現金面交', 'NT$450,000，八德路三段200號'],
        ['115-01-09', '',        'USDT轉帳', '12,610 USDT → 嫌疑人錢包'],
        ['115-01-09', '',        '發現異常', '欲出金遭拒，察覺受騙'],
        ['115-01-11', '08:08',   '製作筆錄', '臺北市松山區派出所，警員吳家宵'],
    ],
}


# ─────────────────────────────────────────────────────────────────────────────
# CLI 入口
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import sys
    mode = sys.argv[1] if len(sys.argv) > 1 else 'case'

    out_dir = os.path.dirname(os.path.dirname(__file__))  # crypto_analyzer/

    if mode == 'template':
        path = os.path.join(out_dir, '虛擬貨幣詐欺案件分析範本_空白.docx')
        build_case_doc({}, path)
        print(f'空白範本已生成：{path}')
    else:
        path = os.path.join(out_dir, '案件分析_陳姿穎_20250111.docx')
        build_case_doc(CASE_1_3_DATA, path)
        print(f'案件分析報告已生成：{path}')
